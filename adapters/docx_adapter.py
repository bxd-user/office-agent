from __future__ import annotations

import re
from pathlib import Path
from typing import Any, Dict, Iterable, List, Set

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

from tools.base import ToolValidationError
from tools.common_tools import ensure_allowed_extension, ensure_file_exists, ensure_parent_dir


SUPPORTED_DOCX_EXTENSIONS = (".docx",)
PLACEHOLDER_PATTERN = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")


# =========================
# 基础文件检查 / 加载 / 保存
# =========================

def ensure_docx_file(file_path: str) -> None:
    try:
        ensure_file_exists(file_path)
    except ToolValidationError as exc:
        if str(exc).startswith("File not found:"):
            raise ToolValidationError(f"Word file not found: {file_path}") from exc
        raise

    ext = Path(file_path).suffix.lower()
    try:
        ensure_allowed_extension(file_path, SUPPORTED_DOCX_EXTENSIONS)
    except ToolValidationError as exc:
        raise ToolValidationError(
            f"Unsupported Word file extension: {ext}. "
            f"Expected one of: {SUPPORTED_DOCX_EXTENSIONS}"
        ) from exc


def load_document_safe(file_path: str) -> DocxDocument:
    ensure_docx_file(file_path)
    try:
        return Document(file_path)
    except Exception as exc:  # noqa: BLE001
        raise ToolValidationError(f"Failed to load Word document: {exc}") from exc


def save_document(doc: DocxDocument, output_path: str) -> str:
    ensure_parent_dir(output_path)
    output = Path(output_path)

    try:
        doc.save(str(output))
    except Exception as exc:  # noqa: BLE001
        raise ToolValidationError(f"Failed to save Word document: {exc}") from exc

    return str(output)


# =========================
# 遍历段落 / 表格
# =========================

def iter_document_paragraphs(doc: DocxDocument) -> Iterable[Paragraph]:
    """遍历文档顶层段落。"""
    for paragraph in doc.paragraphs:
        yield paragraph


def iter_document_tables(doc: DocxDocument) -> Iterable[Table]:
    """遍历文档顶层表格。"""
    for table in doc.tables:
        yield table


def iter_table_cells(table: Table) -> Iterable[_Cell]:
    """遍历表格中的所有单元格。"""
    for row in table.rows:
        for cell in row.cells:
            yield cell


def iter_cell_paragraphs(cell: _Cell) -> Iterable[Paragraph]:
    """遍历单元格中的段落。"""
    for paragraph in cell.paragraphs:
        yield paragraph


def iter_all_paragraphs(doc: DocxDocument) -> Iterable[Paragraph]:
    """统一遍历文档中可见的正文段落 + 表格单元格段落。"""
    for paragraph in iter_document_paragraphs(doc):
        yield paragraph

    for table in iter_document_tables(doc):
        for cell in iter_table_cells(table):
            for paragraph in iter_cell_paragraphs(cell):
                yield paragraph


# =========================
# 文本读取
# =========================

def get_all_paragraph_texts(file_path: str) -> List[str]:
    doc = load_document_safe(file_path)
    return [p.text for p in iter_document_paragraphs(doc)]


def get_all_table_texts(file_path: str) -> List[List[List[str]]]:
    """返回所有表格文本。

    结构：
    [
        [   # table 1
            ["单元格11", "单元格12"],
            ["单元格21", "单元格22"],
        ],
        ...
    ]
    """
    doc = load_document_safe(file_path)
    all_tables: List[List[List[str]]] = []

    for table in iter_document_tables(doc):
        all_tables.append(_table_to_text_rows(table))

    return all_tables


# =========================
# 占位符提取
# =========================

def extract_placeholders_from_text(text: str) -> List[str]:
    return [match.group(1).strip() for match in PLACEHOLDER_PATTERN.finditer(text)]


def extract_placeholders_from_document(doc: DocxDocument) -> List[str]:
    """提取文档对象中的占位符名称，不重复，保持出现顺序。"""
    found: List[str] = []
    seen: Set[str] = set()

    for paragraph in iter_all_paragraphs(doc):
        for key in extract_placeholders_from_text(paragraph.text):
            if key not in seen:
                seen.add(key)
                found.append(key)

    return found


def extract_placeholders(file_path: str) -> List[str]:
    """提取文档中所有占位符名称，不重复，保持出现顺序。"""
    doc = load_document_safe(file_path)
    return extract_placeholders_from_document(doc)


def find_unfilled_placeholders(file_path: str) -> List[str]:
    """查找当前文档里仍然存在的占位符。"""
    return extract_placeholders(file_path)


# =========================
# 替换逻辑
# =========================

def build_placeholder_token(key: str) -> str:
    return f"{{{{{key}}}}}"


def normalize_replacements(replacements: Dict[str, Any]) -> Dict[str, str]:
    """把 replacements 标准化为字符串字典。"""
    normalized: Dict[str, str] = {}
    for key, value in replacements.items():
        normalized[str(key).strip()] = "" if value is None else str(value)
    return normalized


def replace_placeholders_in_text(text: str, replacements: Dict[str, str]) -> str:
    """替换文本中的 {{字段}}。

    支持：
    - {{姓名}}
    - {{ 姓名 }}
    """
    result = text
    for key, value in replacements.items():
        pattern = re.compile(r"\{\{\s*" + re.escape(key) + r"\s*\}\}")
        result = pattern.sub(value, result)
    return result


def replace_paragraph_text(paragraph: Paragraph, replacements: Dict[str, str]) -> bool:
    """替换一个段落里的占位符。

    这里采用“段落整段替换”策略：
    - 优点：对段落/表格都稳定
    - 缺点：如果该段落内 run 样式很多，局部样式可能丢失

    对你当前 demo，这个取舍是合理的。
    """
    original = paragraph.text
    new_text = replace_placeholders_in_text(original, replacements)

    if new_text == original:
        return False

    paragraph.text = new_text
    return True


def replace_placeholders_in_doc(
    file_path: str,
    replacements: Dict[str, Any],
    output_path: str,
) -> str:
    """替换文档中的占位符并输出新文件。"""
    doc = load_document_safe(file_path)
    normalized = normalize_replacements(replacements)

    for paragraph in iter_all_paragraphs(doc):
        replace_paragraph_text(paragraph, normalized)

    return save_document(doc, output_path)


# =========================
# 文档检查 / 结构摘要
# =========================

def inspect_document(file_path: str) -> Dict[str, Any]:
    doc = load_document_safe(file_path)

    paragraph_texts = [p.text for p in iter_document_paragraphs(doc)]
    tables = list(iter_document_tables(doc))
    placeholders = extract_placeholders_from_document(doc)

    table_summaries = []
    for idx, table in enumerate(tables, start=1):
        row_count = len(table.rows)
        col_count = max((len(row.cells) for row in table.rows), default=0)
        table_summaries.append(
            {
                "table_index": idx,
                "row_count": row_count,
                "column_count": col_count,
            }
        )

    non_empty_paragraphs = [t for t in paragraph_texts if t.strip()]

    return {
        "file_path": file_path,
        "paragraph_count": len(paragraph_texts),
        "non_empty_paragraph_count": len(non_empty_paragraphs),
        "table_count": len(tables),
        "tables": table_summaries,
        "placeholders": placeholders,
        "preview_paragraphs": non_empty_paragraphs[:10],
    }


# =========================
# 额外辅助：结构化获取段落/表格内容
# =========================

def get_document_blocks(file_path: str) -> List[Dict[str, Any]]:
    """把文档按 block 形式返回，便于后续 inspector/tool 使用。

    返回示例：
    [
        {"type": "paragraph", "text": "标题"},
        {"type": "paragraph", "text": "姓名：张三"},
        {"type": "table", "rows": [["A", "B"], ["C", "D"]]},
    ]
    """
    doc = load_document_safe(file_path)
    blocks: List[Dict[str, Any]] = []

    # 先返回顶层段落
    for paragraph in doc.paragraphs:
        blocks.append(
            {
                "type": "paragraph",
                "text": paragraph.text,
            }
        )

    # 再返回顶层表格
    for table in doc.tables:
        blocks.append(
            {
                "type": "table",
                "rows": _table_to_text_rows(table),
            }
        )

    return blocks


def _cell_text(cell: _Cell) -> str:
    return "\n".join(p.text for p in cell.paragraphs).strip()


def _table_to_text_rows(table: Table) -> List[List[str]]:
    rows: List[List[str]] = []
    for row in table.rows:
        rows.append([_cell_text(cell) for cell in row.cells])
    return rows