from pathlib import Path

from openpyxl import Workbook
from docx import Document


def build_excel_xlsx(path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "人员信息"
    ws.append(["姓名", "性别", "出生年月", "民族", "籍贯", "工作职务", "申请入党时间", "团员编号", "团支部", "被推荐人优缺点", "团支部意见"])
    ws.append(["张三", "男", "1998-05", "汉族", "北京", "工程师", "2023-06", "TY2026001", "软件一支部", "技术扎实，责任心强", "同意推荐"])
    ws.append(["李四", "女", "1999-09", "汉族", "上海", "产品经理", "2024-01", "TY2026002", "产品二支部", "沟通能力强，执行力好", "建议重点培养"])

    ws2 = wb.create_sheet("补充信息")
    ws2.append(["姓名", "邮箱", "联系电话"])
    ws2.append(["张三", "zhangsan@example.com", "13800000001"])
    ws2.append(["李四", "lisi@example.com", "13800000002"])

    wb.save(path)


def build_excel_csv(path: Path) -> None:
    content = "姓名,性别,出生年月,民族,籍贯,工作职务\n张三,男,1998-05,汉族,北京,工程师\n李四,女,1999-09,汉族,上海,产品经理\n"
    path.write_text(content, encoding="utf-8-sig")


def build_word_template(path: Path) -> None:
    doc = Document()
    doc.add_heading("发展对象推荐表（测试样例）", level=1)
    doc.add_paragraph("姓名：{{姓名}}")
    doc.add_paragraph("性别：${性别}")
    doc.add_paragraph("出生年月：【出生年月】")
    doc.add_paragraph("民族：[[民族]]")
    doc.add_paragraph("籍贯：____")

    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "团支部"
    table.cell(0, 1).text = ""
    table.cell(1, 0).text = "团员编号"
    table.cell(1, 1).text = ""
    table.cell(2, 0).text = "工作职务"
    table.cell(2, 1).text = ""
    table.cell(3, 0).text = "申请入党时间"
    table.cell(3, 1).text = ""
    table.cell(4, 0).text = "被推荐人优缺点"
    table.cell(4, 1).text = ""
    table.cell(5, 0).text = "团支部意见"
    table.cell(5, 1).text = ""

    doc.save(path)


def main() -> None:
    sample_dir = Path("storage/uploads/samples")
    sample_dir.mkdir(parents=True, exist_ok=True)

    xlsx_path = sample_dir / "sample_people.xlsx"
    csv_path = sample_dir / "sample_people.csv"
    word_path = sample_dir / "sample_review_template.docx"

    build_excel_xlsx(xlsx_path)
    build_excel_csv(csv_path)
    build_word_template(word_path)

    print("生成完成:")
    print(xlsx_path)
    print(csv_path)
    print(word_path)


if __name__ == "__main__":
    main()
