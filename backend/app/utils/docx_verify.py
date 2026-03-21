from docx import Document


def count_unreplaced_placeholders(file_path: str) -> int:
    doc = Document(file_path)
    count = 0

    for para in doc.paragraphs:
        if "{{" in para.text and "}}" in para.text:
            count += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "{{" in cell.text and "}}" in cell.text:
                    count += 1

    return count