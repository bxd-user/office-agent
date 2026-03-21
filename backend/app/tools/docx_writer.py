from pathlib import Path
from docx import Document


class DocxWriter:
    def replace_placeholders(
        self,
        template_path: str,
        output_path: str,
        data: dict[str, str],
    ) -> str:
        doc = Document(template_path)

        # 替换段落中的占位符
        for para in doc.paragraphs:
            for key, value in data.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in para.text:
                    para.text = para.text.replace(placeholder, str(value))

        # 替换表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in data.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(value))

        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_file))

        return str(output_file)