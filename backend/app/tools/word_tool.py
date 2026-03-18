import os
import re
from dataclasses import dataclass, asdict
from typing import Optional

from docx import Document
from app.tools.base_tool import BaseDataTool


@dataclass
class FieldSlot:
    field_name: str
    normalized_name: str
    source_type: str  # paragraph_label / table_label
    location: dict
    target: Optional[dict]
    confidence: float


class WordTool(BaseDataTool):
    TOOL_NAME = "word"
    SUPPORTED_WORD_EXTENSIONS = {".docx", ".docm", ".dotx", ".dotm"}
    SUPPORTED_EXTENSIONS = SUPPORTED_WORD_EXTENSIONS
    PLACEHOLDER_PATTERN = re.compile(
        r"\{\{\s*([^{}]+?)\s*\}\}|\$\{\s*([^{}]+?)\s*\}|【\s*([^【】]+?)\s*】|\[\[\s*([^\[\]]+?)\s*\]\]"
    )
    LABEL_ENDINGS = ("：", ":")
    INLINE_LABEL_PATTERN = re.compile(r"^\s*([^:：]{1,80})\s*[:：]\s*(.*?)\s*$")

    def read_for_llm(self, path: str, **kwargs) -> dict:
        return self.read_document_for_llm(
            path=path,
            max_paragraphs_per_area=int(kwargs.get("max_paragraphs_per_area", 3000)),
            max_tables_per_area=int(kwargs.get("max_tables_per_area", 300)),
            max_rows_per_table=int(kwargs.get("max_rows_per_table", 500)),
            max_cols_per_row=int(kwargs.get("max_cols_per_row", 100)),
        )

    def write_from_llm(self, source_path: str, data: dict, output_path: str, **kwargs) -> str:
        return self.write_document_from_llm(source_path=source_path, document_data=data, output_path=output_path)

    def execute_llm_instruction(self, instruction: dict) -> dict:
        action = str(instruction.get("action") or "").strip().lower()

        if action == "read":
            path = str(instruction.get("path") or "")
            options = instruction.get("options", {}) if isinstance(instruction.get("options", {}), dict) else {}
            return {
                "action": "read",
                "data": self.read_for_llm(path, **options),
            }

        if action == "write":
            source_path = str(instruction.get("source_path") or "")
            output_path = str(instruction.get("output_path") or "")
            data = instruction.get("data", {}) if isinstance(instruction.get("data", {}), dict) else {}
            return {
                "action": "write",
                "output_path": self.write_from_llm(source_path=source_path, data=data, output_path=output_path),
            }

        if action == "fill_context":
            template_path = str(instruction.get("template_path") or "")
            output_path = str(instruction.get("output_path") or "")
            context = instruction.get("context", {}) if isinstance(instruction.get("context", {}), dict) else {}
            return {
                "action": "fill_context",
                "output_path": self.fill_document_with_context(
                    template_path=template_path,
                    context=context,
                    output_path=output_path,
                ),
            }

        if action == "extract_fields":
            path = str(instruction.get("path") or "")
            data = self.read_document_for_llm(path)
            return {
                "action": "extract_fields",
                "fields": data.get("fields", []),
                "field_slots": data.get("field_slots", []),
            }

        raise ValueError(f"WordTool 不支持的 action: {action}")

    def get_supported_form_fields(self) -> list[str]:
        return []

    def get_supported_form_fields_from_template(self, template_path: str) -> list[str]:
        payload = self.read_document_for_llm(template_path)
        return payload.get("fields", [])

    def read_document_for_llm(
        self,
        path: str,
        max_paragraphs_per_area: int = 3000,
        max_tables_per_area: int = 300,
        max_rows_per_table: int = 500,
        max_cols_per_row: int = 100,
    ) -> dict:
        self._ensure_supported_word_file(path)
        doc = Document(path)

        body = self._extract_area_content(
            paragraphs=doc.paragraphs,
            tables=doc.tables,
            max_paragraphs=max_paragraphs_per_area,
            max_tables=max_tables_per_area,
            max_rows=max_rows_per_table,
            max_cols=max_cols_per_row,
        )

        headers = []
        footers = []

        for section_index, section in enumerate(doc.sections):
            headers.append(
                {
                    "section_index": section_index,
                    **self._extract_area_content(
                        paragraphs=section.header.paragraphs,
                        tables=section.header.tables,
                        max_paragraphs=max_paragraphs_per_area,
                        max_tables=max_tables_per_area,
                        max_rows=max_rows_per_table,
                        max_cols=max_cols_per_row,
                    ),
                }
            )
            footers.append(
                {
                    "section_index": section_index,
                    **self._extract_area_content(
                        paragraphs=section.footer.paragraphs,
                        tables=section.footer.tables,
                        max_paragraphs=max_paragraphs_per_area,
                        max_tables=max_tables_per_area,
                        max_rows=max_rows_per_table,
                        max_cols=max_cols_per_row,
                    ),
                }
            )

        field_slots = self._inspect_doc_field_slots(doc)

        fields = []
        seen = set()
        for slot in field_slots:
            normalized = slot.normalized_name
            if normalized and normalized not in seen:
                seen.add(normalized)
                fields.append(slot.field_name)

        return {
            "file_type": os.path.splitext(path)[1].lower(),
            "fields": fields,
            "field_slots": [asdict(slot) for slot in field_slots],
            "stats": {
                "body_paragraphs": len(body.get("paragraphs", [])),
                "body_tables": len(body.get("tables", [])),
                "sections": len(doc.sections),
            },
            "content": {
                "body": body,
                "headers": headers,
                "footers": footers,
            },
        }

    def write_document_from_llm(self, source_path: str, document_data: dict, output_path: str) -> str:
        self._ensure_supported_word_file(source_path)
        self._ensure_supported_word_file(output_path)

        doc = Document(source_path)
        content = document_data.get("content", {}) if isinstance(document_data, dict) else {}

        self._apply_area_content(
            paragraphs=doc.paragraphs,
            tables=doc.tables,
            area_data=content.get("body", {}),
        )

        headers_data = content.get("headers", []) if isinstance(content.get("headers", []), list) else []
        footers_data = content.get("footers", []) if isinstance(content.get("footers", []), list) else []

        for header_item in headers_data:
            section_index = self._safe_int(header_item.get("section_index"), default=-1)
            if section_index < 0 or section_index >= len(doc.sections):
                continue
            self._apply_area_content(
                paragraphs=doc.sections[section_index].header.paragraphs,
                tables=doc.sections[section_index].header.tables,
                area_data=header_item,
            )

        for footer_item in footers_data:
            section_index = self._safe_int(footer_item.get("section_index"), default=-1)
            if section_index < 0 or section_index >= len(doc.sections):
                continue
            self._apply_area_content(
                paragraphs=doc.sections[section_index].footer.paragraphs,
                tables=doc.sections[section_index].footer.tables,
                area_data=footer_item,
            )

        doc.save(output_path)
        return output_path

    def fill_review_form(self, template_path: str, context: dict, output_path: str) -> str:
        return self.fill_document_with_context(template_path, context, output_path)

    def fill_document_with_context(self, template_path: str, context: dict, output_path: str) -> str:
        self._ensure_supported_word_file(template_path)
        doc = Document(template_path)
        normalized_context = self._build_normalized_context(context)

        self._fill_doc_with_context(doc, normalized_context)

        doc.save(output_path)
        return output_path

    def _fill_doc_with_context(self, doc, normalized_context: dict) -> None:
        # 先处理正文中的普通段落/占位符
        self._fill_area_with_context(doc.paragraphs, doc.tables, normalized_context)

        # 再处理正文表格中的字段槽位
        self._fill_doc_with_field_slots(doc, normalized_context)

        # 处理页眉页脚
        for section_index, section in enumerate(doc.sections):
            self._fill_area_with_context(section.header.paragraphs, section.header.tables, normalized_context)
            self._fill_area_with_context(section.footer.paragraphs, section.footer.tables, normalized_context)

            self._fill_area_tables_with_field_slots(
                section.header.tables,
                normalized_context,
                area="header",
                section_index=section_index,
            )
            self._fill_area_tables_with_field_slots(
                section.footer.tables,
                normalized_context,
                area="footer",
                section_index=section_index,
            )

    def _fill_area_with_context(self, paragraphs, tables, normalized_context: dict) -> None:
        for paragraph in paragraphs:
            self._fill_paragraph(paragraph, normalized_context)
        for table in tables:
            self._fill_table_placeholders(table, normalized_context)
            self._fill_table_by_label_position_auto(table, normalized_context)

    def _fill_doc_with_field_slots(self, doc, normalized_context: dict) -> None:
        self._fill_area_tables_with_field_slots(doc.tables, normalized_context, area="body", section_index=None)

    def _fill_area_tables_with_field_slots(self, tables, normalized_context: dict, area: str, section_index: Optional[int]) -> None:
        for table_index, table in enumerate(tables):
            slots = self._extract_table_field_slots(table, table_index=table_index, area=area, section_index=section_index)
            self._apply_field_slots_to_table(table, slots, normalized_context)

    def _extract_area_content(
        self,
        paragraphs,
        tables,
        max_paragraphs: int,
        max_tables: int,
        max_rows: int,
        max_cols: int,
    ) -> dict:
        paragraph_items = []
        for index, paragraph in enumerate(paragraphs):
            if index >= max_paragraphs:
                break
            paragraph_items.append({"index": index, "text": paragraph.text or ""})

        table_items = []
        for table_index, table in enumerate(tables):
            if table_index >= max_tables:
                break
            rows = []
            for row_index, row in enumerate(table.rows):
                if row_index >= max_rows:
                    break
                cells = []
                for col_index, cell in enumerate(row.cells):
                    if col_index >= max_cols:
                        break
                    cells.append(cell.text or "")
                rows.append(cells)
            table_items.append({"index": table_index, "rows": rows})

        return {"paragraphs": paragraph_items, "tables": table_items}

    def _apply_area_content(self, paragraphs, tables, area_data: dict) -> None:
        paragraph_data = area_data.get("paragraphs", []) if isinstance(area_data, dict) else []
        table_data = area_data.get("tables", []) if isinstance(area_data, dict) else []

        self._apply_paragraph_data(paragraphs, paragraph_data)
        self._apply_table_data(tables, table_data)

    def _apply_paragraph_data(self, paragraphs, paragraph_data) -> None:
        if not isinstance(paragraph_data, list):
            return

        for item_index, item in enumerate(paragraph_data):
            if isinstance(item, str):
                paragraph_index = item_index
                text = item
            elif isinstance(item, dict):
                paragraph_index = self._safe_int(item.get("index"), default=item_index)
                text = "" if item.get("text") is None else str(item.get("text"))
            else:
                continue

            if 0 <= paragraph_index < len(paragraphs):
                paragraphs[paragraph_index].text = text

    def _apply_table_data(self, tables, table_data) -> None:
        if not isinstance(table_data, list):
            return

        for item_index, table_item in enumerate(table_data):
            if not isinstance(table_item, dict):
                continue

            table_index = self._safe_int(table_item.get("index"), default=item_index)
            if table_index < 0 or table_index >= len(tables):
                continue

            rows = table_item.get("rows", [])
            if not isinstance(rows, list):
                continue

            self._replace_table_values(tables[table_index], rows)

    def _replace_table_values(self, table, rows) -> None:
        existing_rows = len(table.rows)
        target_rows = len(rows)

        while len(table.rows) < target_rows:
            table.add_row()

        for row_index, row_data in enumerate(rows):
            if not isinstance(row_data, list):
                continue

            if row_index >= len(table.rows):
                break

            table_row = table.rows[row_index]
            limit = min(len(table_row.cells), len(row_data))
            for col_index in range(limit):
                value = "" if row_data[col_index] is None else str(row_data[col_index])
                table_row.cells[col_index].text = value

            if len(row_data) > len(table_row.cells) and len(table_row.cells) > 0:
                overflow = row_data[len(table_row.cells):]
                tail = "\n".join(str(v) for v in overflow if v is not None)
                if tail:
                    last_cell = table_row.cells[-1]
                    joiner = "\n" if (last_cell.text or "").strip() else ""
                    last_cell.text = f"{last_cell.text}{joiner}{tail}"

        if target_rows < existing_rows:
            for row_index in range(target_rows, existing_rows):
                for cell in table.rows[row_index].cells:
                    cell.text = ""

    def _inspect_doc_field_slots(self, doc) -> list[FieldSlot]:
        slots: list[FieldSlot] = []

        # 段落字段
        for p_index, paragraph in enumerate(doc.paragraphs):
            slots.extend(self._extract_paragraph_field_slots(paragraph, paragraph_index=p_index, area="body", section_index=None))

        # 正文表格字段
        for table_index, table in enumerate(doc.tables):
            slots.extend(self._extract_table_field_slots(table, table_index=table_index, area="body", section_index=None))

        # 页眉页脚
        for section_index, section in enumerate(doc.sections):
            for p_index, paragraph in enumerate(section.header.paragraphs):
                slots.extend(
                    self._extract_paragraph_field_slots(
                        paragraph,
                        paragraph_index=p_index,
                        area="header",
                        section_index=section_index,
                    )
                )
            for table_index, table in enumerate(section.header.tables):
                slots.extend(
                    self._extract_table_field_slots(
                        table,
                        table_index=table_index,
                        area="header",
                        section_index=section_index,
                    )
                )

            for p_index, paragraph in enumerate(section.footer.paragraphs):
                slots.extend(
                    self._extract_paragraph_field_slots(
                        paragraph,
                        paragraph_index=p_index,
                        area="footer",
                        section_index=section_index,
                    )
                )
            for table_index, table in enumerate(section.footer.tables):
                slots.extend(
                    self._extract_table_field_slots(
                        table,
                        table_index=table_index,
                        area="footer",
                        section_index=section_index,
                    )
                )

        # 去重
        deduped = []
        seen = set()
        for slot in sorted(slots, key=lambda x: (-x.confidence, x.field_name)):
            key = (
                slot.normalized_name,
                slot.source_type,
                tuple(sorted(slot.location.items())) if isinstance(slot.location, dict) else str(slot.location),
            )
            if key in seen:
                continue
            seen.add(key)
            deduped.append(slot)

        return deduped

    def _extract_paragraph_field_slots(self, paragraph, paragraph_index: int, area: str, section_index: Optional[int]) -> list[FieldSlot]:
        slots: list[FieldSlot] = []
        text = (paragraph.text or "").strip()
        if not text:
            return slots

        # 占位符
        for match in self.PLACEHOLDER_PATTERN.finditer(text):
            field_name = self._match_field_name(match)
            normalized = self._normalize(field_name)
            if not normalized:
                continue
            slots.append(
                FieldSlot(
                    field_name=field_name,
                    normalized_name=normalized,
                    source_type="paragraph_placeholder",
                    location={
                        "area": area,
                        "section_index": section_index,
                        "paragraph_index": paragraph_index,
                    },
                    target={
                        "kind": "paragraph",
                        "area": area,
                        "section_index": section_index,
                        "paragraph_index": paragraph_index,
                    },
                    confidence=0.99,
                )
            )

        # 冒号标签，如 团支部：
        label = self._extract_label_name(text)
        if label:
            normalized = self._normalize(label)
            slots.append(
                FieldSlot(
                    field_name=label,
                    normalized_name=normalized,
                    source_type="paragraph_label",
                    location={
                        "area": area,
                        "section_index": section_index,
                        "paragraph_index": paragraph_index,
                    },
                    target={
                        "kind": "paragraph_inline",
                        "area": area,
                        "section_index": section_index,
                        "paragraph_index": paragraph_index,
                    },
                    confidence=0.96,
                )
            )
            return slots

        # 行内标签，如 团支部：____
        inline_match = self.INLINE_LABEL_PATTERN.match(text)
        if inline_match:
            label = self._clean_field_name(inline_match.group(1))
            normalized = self._normalize(label)
            if normalized:
                slots.append(
                    FieldSlot(
                        field_name=label,
                        normalized_name=normalized,
                        source_type="paragraph_inline_label",
                        location={
                            "area": area,
                            "section_index": section_index,
                            "paragraph_index": paragraph_index,
                        },
                        target={
                            "kind": "paragraph_inline",
                            "area": area,
                            "section_index": section_index,
                            "paragraph_index": paragraph_index,
                        },
                        confidence=0.95,
                    )
                )

        return slots

    def _extract_table_field_slots(self, table, table_index: int, area: str, section_index: Optional[int]) -> list[FieldSlot]:
        slots: list[FieldSlot] = []
        rows = table.rows
        row_count = len(rows)

        for row_index, row in enumerate(rows):
            col_count = len(row.cells)
            for col_index, cell in enumerate(row.cells):
                raw_text = cell.text or ""
                label_name = self._extract_label_name(raw_text)

                if not label_name:
                    inline_match = self.INLINE_LABEL_PATTERN.match(raw_text.strip())
                    if inline_match:
                        label_name = self._clean_field_name(inline_match.group(1))

                # 新增：表格短标签识别（不依赖冒号）
                if not label_name and self._is_table_label_cell(raw_text):
                    label_name = self._clean_field_name(raw_text)

                if not label_name:
                    continue

                normalized = self._normalize(label_name)
                if not normalized:
                    continue

                target = self._find_best_target_cell(rows, row_index, col_index)
                confidence = 0.65
                if target is not None:
                    if target["row"] == row_index and target["col"] == col_index + 1:
                        confidence = 0.96
                    elif target["row"] == row_index:
                        confidence = 0.88
                    elif target["col"] == col_index:
                        confidence = 0.80
                    else:
                        confidence = 0.72

                slots.append(
                    FieldSlot(
                        field_name=label_name,
                        normalized_name=normalized,
                        source_type="table_label",
                        location={
                            "area": area,
                            "section_index": section_index,
                            "table_index": table_index,
                            "row": row_index,
                            "col": col_index,
                        },
                        target=None if target is None else {
                            "kind": "table_cell",
                            "area": area,
                            "section_index": section_index,
                            "table_index": table_index,
                            "row": target["row"],
                            "col": target["col"],
                        },
                        confidence=confidence,
                    )
                )

        return slots

    def _apply_field_slots_to_table(self, table, slots: list[FieldSlot], normalized_context: dict) -> None:
        for slot in slots:
            value = normalized_context.get(slot.normalized_name, "")
            if value == "":
                continue
            if not slot.target or slot.target.get("kind") != "table_cell":
                continue

            row_index = self._safe_int(slot.target.get("row"), default=-1)
            col_index = self._safe_int(slot.target.get("col"), default=-1)

            if row_index < 0 or row_index >= len(table.rows):
                continue
            row = table.rows[row_index]
            if col_index < 0 or col_index >= len(row.cells):
                continue

            target_cell = row.cells[col_index]
            current_text = target_cell.text or ""

            if self._can_write_to_cell(current_text):
                target_cell.text = value

    def _find_best_target_cell(self, rows, row_index: int, col_index: int) -> Optional[dict]:
        row_count = len(rows)
        col_count = len(rows[row_index].cells) if row_index < row_count else 0

        candidates = []

        # 1. 同行右侧 1~3 格，优先级最高
        for offset in range(1, 4):
            c = col_index + offset
            if c >= col_count:
                break
            cell = rows[row_index].cells[c]
            score = self._score_target_cell(cell.text or "", same_row=True, distance=offset)
            if score > 0:
                candidates.append({"row": row_index, "col": c, "score": score})

        # 2. 下一行同列 / 右侧 1 格
        for offset in range(1, 3):
            r = row_index + offset
            if r >= row_count:
                break

            if col_index < len(rows[r].cells):
                cell = rows[r].cells[col_index]
                score = self._score_target_cell(cell.text or "", same_row=False, distance=offset)
                if score > 0:
                    candidates.append({"row": r, "col": col_index, "score": score})

            if col_index + 1 < len(rows[r].cells):
                cell = rows[r].cells[col_index + 1]
                score = self._score_target_cell(cell.text or "", same_row=False, distance=offset + 0.5)
                if score > 0:
                    candidates.append({"row": r, "col": col_index + 1, "score": score})

        if not candidates:
            return None

        candidates.sort(key=lambda x: x["score"], reverse=True)
        return candidates[0]

    def _score_target_cell(self, text: str, same_row: bool, distance: float) -> float:
        value = text or ""

        # 看起来像另一个标签，直接降到很低
        if self._looks_like_label(value):
            return -1.0

        score = 0.0

        if self._can_write_to_cell(value):
            score += 1.2
        elif len(self._normalize(value)) <= 2:
            score += 0.2
        else:
            score -= 0.6

        if same_row:
            score += 0.5

        score += max(0.0, 0.4 - 0.1 * float(distance))
        return score

    def _can_write_to_cell(self, text: str) -> bool:
        value = text or ""
        normalized = self._normalize(value)
        if normalized == "":
            return True
        return self._is_placeholder_like(value)

    def _extract_fields_from_texts(self, texts: list[str]) -> list[str]:
        fields = []
        seen = set()

        def add_field(raw: str) -> None:
            value = self._clean_field_name(raw)
            normalized = self._normalize(value)
            if not normalized or normalized in seen:
                return
            seen.add(normalized)
            fields.append(value)

        for text in texts:
            value = text or ""
            for match in self.PLACEHOLDER_PATTERN.finditer(value):
                add_field(self._match_field_name(match))

            label = self._extract_label_name(value)
            if label:
                add_field(label)

            inline_match = self.INLINE_LABEL_PATTERN.match(value.strip())
            if inline_match:
                add_field(inline_match.group(1))

        return fields

    def _ensure_supported_word_file(self, path: str) -> None:
        ext = os.path.splitext(path)[1].lower()
        if ext not in self.SUPPORTED_WORD_EXTENSIONS:
            raise ValueError(
                f"当前支持的 Word 模板类型: {sorted(self.SUPPORTED_WORD_EXTENSIONS)}，收到: {ext or '无扩展名'}"
            )

    def _normalize(self, text: str) -> str:
        if text is None:
            return ""
        text = str(text)
        text = text.replace("\u3000", " ")
        text = re.sub(r"[\s\r\n\t]+", "", text)
        text = re.sub(r"[：:]", "", text)
        return text

    def _safe_int(self, value, default: int = 0) -> int:
        try:
            return int(value)
        except (TypeError, ValueError):
            return default

    def _clean_field_name(self, raw: str) -> str:
        text = str(raw or "").strip()
        text = re.sub(r"\s+", " ", text)
        for ending in self.LABEL_ENDINGS:
            if text.endswith(ending):
                text = text[: -len(ending)].strip()
        return text

    def _extract_label_name(self, text: str) -> str:
        value = str(text or "").strip()
        if not value:
            return ""

        for ending in self.LABEL_ENDINGS:
            if ending in value:
                head, tail = value.split(ending, 1)
                if tail.strip():
                    continue
                candidate = self._clean_field_name(head)
                if 1 <= len(candidate) <= 80:
                    return candidate

        return ""

    def _match_field_name(self, match) -> str:
        return self._clean_field_name(match.group(1) or match.group(2) or match.group(3) or match.group(4) or "")

    def _build_normalized_context(self, context: dict) -> dict:
        normalized = {}
        for key, value in context.items():
            normalized_key = self._normalize(key)
            if not normalized_key:
                continue
            normalized[normalized_key] = "" if value is None else str(value)
        return normalized

    def _get_context_value(self, normalized_context: dict, key: str) -> str:
        return normalized_context.get(self._normalize(key), "")

    def _fill_paragraph(self, paragraph, normalized_context: dict) -> None:
        original = paragraph.text or ""
        if not original:
            return

        self._replace_placeholders_in_runs(paragraph, normalized_context)

        after_runs = paragraph.text or ""
        if after_runs != original:
            original = after_runs

        replaced = self._replace_placeholders_in_text(original, normalized_context)
        if replaced != original:
            paragraph.text = replaced
            return

        filled = self._fill_inline_label_text(original, normalized_context)
        if filled is not None and filled != original:
            paragraph.text = filled

    def _fill_table_placeholders(self, table, normalized_context: dict) -> None:
        for row in table.rows:
            for cell in row.cells:
                original = cell.text or ""
                replaced = self._replace_placeholders_in_text(original, normalized_context)
                if replaced != original:
                    cell.text = replaced

    def _replace_placeholders_in_runs(self, paragraph, normalized_context: dict) -> None:
        for run in paragraph.runs:
            text = run.text or ""
            if not text:
                continue
            replaced = self._replace_placeholders_in_text(text, normalized_context)
            if replaced != text:
                run.text = replaced

    def _replace_placeholders_in_text(self, text: str, normalized_context: dict) -> str:
        def replacement(match):
            field_name = self._match_field_name(match)
            value = self._get_context_value(normalized_context, field_name)
            return value if value != "" else match.group(0)

        return self.PLACEHOLDER_PATTERN.sub(replacement, text)

    def _fill_inline_label_text(self, text: str, normalized_context: dict) -> str | None:
        stripped = (text or "").strip()
        if not stripped:
            return None

        match = self.INLINE_LABEL_PATTERN.match(stripped)
        if not match:
            return None

        label = self._clean_field_name(match.group(1))
        right_text = (match.group(2) or "").strip()
        value = self._get_context_value(normalized_context, label)
        if value == "":
            return None

        if right_text and not self._is_placeholder_like(right_text):
            return None

        separator = "：" if "：" in stripped else ":"
        return f"{label}{separator}{value}"

    def _is_placeholder_like(self, text: str) -> bool:
        if not text:
            return True
        normalized = re.sub(r"[\s_＿\-—\.·•。:：()（）]+", "", text)
        return normalized == ""

    def _is_table_label_cell(self, text: str) -> bool:
        value = (text or "").strip()
        if not value:
            return False

        normalized = self._normalize(value)
        if not normalized:
            return False

        # 太长通常不是标签
        if len(normalized) > 20:
            return False

        # 纯数字/日期/编号通常不是标签
        if any(ch.isdigit() for ch in normalized):
            return False

        # 含明显长句标点，通常不是标签
        if any(p in value for p in ["。", "，", ",", ";", "；"]):
            return False

        # 有占位符，说明本身不是标签
        if self.PLACEHOLDER_PATTERN.search(value):
            return False

        # 像标签的短文本
        return True

    def _looks_like_label(self, text: str) -> bool:
        value = (text or "").strip()
        if not value:
            return False

        if self._extract_label_name(value):
            return True

        if self.INLINE_LABEL_PATTERN.match(value):
            return True

        return self._is_table_label_cell(value)

    def _fill_table_by_label_position_auto(self, table, normalized_context: dict) -> None:
        rows = table.rows

        for row_index, row in enumerate(rows):
            for col_index, cell in enumerate(row.cells):
                raw_text = cell.text or ""
                label_name = self._extract_label_name(raw_text)

                if not label_name:
                    inline_match = self.INLINE_LABEL_PATTERN.match(raw_text.strip())
                    if inline_match:
                        label_name = self._clean_field_name(inline_match.group(1))

                if not label_name and self._is_table_label_cell(raw_text):
                    label_name = self._clean_field_name(raw_text)

                if not label_name:
                    continue

                normalized_field = self._normalize(label_name)
                if normalized_field not in normalized_context:
                    continue

                field_value = normalized_context[normalized_field]
                target = self._find_best_target_cell(rows, row_index, col_index)
                if not target:
                    continue

                target_cell = rows[target["row"]].cells[target["col"]]
                if self._can_write_to_cell(target_cell.text):
                    target_cell.text = field_value