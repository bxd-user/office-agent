from pathlib import Path
from typing import List

from docx import Document

from app.domain.models import ParsedDocument, ParsedTable


class DocxReader:
    """
    负责读取 .docx 文件，并转换为结构化数据
    """

    def read(self, file_path: str) -> ParsedDocument:
        path = Path(file_path)
        doc = Document(file_path)

        # ===== 1. 读取段落 =====
        paragraphs: List[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # ===== 2. 读取表格 =====
        tables: List[ParsedTable] = []
        table_text_blocks: List[str] = []

        for table in doc.tables:
            rows = []

            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                rows.append(row_data)

            tables.append(ParsedTable(rows=rows))

            # 同时转成文本，方便 LLM 使用
            for row in rows:
                line = " | ".join(cell for cell in row if cell)
                if line:
                    table_text_blocks.append(line)

        # ===== 3. 拼接 full_text（非常重要）=====
        full_text_parts = []

        if paragraphs:
            full_text_parts.append("\n".join(paragraphs))

        if table_text_blocks:
            full_text_parts.append("\n".join(table_text_blocks))

        full_text = "\n\n".join(full_text_parts)

        return ParsedDocument(
            file_name=path.name,
            file_path=str(path),
            paragraphs=paragraphs,
            tables=tables,
            full_text=full_text,
        )

    def read_text(self, file_path: str) -> str:
        parsed = self.read(file_path)
        return parsed.full_text
