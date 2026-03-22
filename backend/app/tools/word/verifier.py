import re
from typing import Any, Dict, List

from docx import Document


PLACEHOLDER_PATTERN = re.compile(r"\{\{(.*?)\}\}")


class DocxVerifier:
    def verify_filled_document(
        self,
        output_path: str,
        expected_placeholders: List[str],
        filled_data: Dict[str, Any],
    ) -> Dict[str, Any]:
        doc = Document(output_path)

        unreplaced_placeholders: List[str] = []

        # 检查段落
        for para in doc.paragraphs:
            unreplaced_placeholders.extend(self._find_placeholders(para.text))

        # 检查表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    unreplaced_placeholders.extend(self._find_placeholders(cell.text))

        # 去重
        unreplaced_placeholders = self._dedupe(unreplaced_placeholders)

        empty_fields = []
        for key in expected_placeholders:
            value = filled_data.get(key, "")
            if value is None or str(value).strip() == "":
                empty_fields.append(key)

        verify_passed = not unreplaced_placeholders and not empty_fields

        return {
            "verify_passed": verify_passed,
            "unreplaced_placeholders": unreplaced_placeholders,
            "empty_fields": empty_fields,
            "needs_repair": not verify_passed,
        }

    def _find_placeholders(self, text: str) -> List[str]:
        if not text:
            return []
        return [item.strip() for item in PLACEHOLDER_PATTERN.findall(text) if item.strip()]

    def _dedupe(self, items: List[str]) -> List[str]:
        result = []
        seen = set()
        for item in items:
            if item not in seen:
                seen.add(item)
                result.append(item)
        return result