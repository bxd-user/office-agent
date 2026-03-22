from pathlib import Path
from uuid import uuid4

from docx import Document


class DocxWriter:
    def replace_placeholders(
        self,
        template_path: str,
        output_path: str,
        data: dict[str, str],
    ) -> str:
        doc = Document(template_path)

        # 替换段落中的占位符
        for para in doc.paragraphs:
            for key, value in data.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in para.text:
                    para.text = para.text.replace(placeholder, str(value))

        # 替换表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in data.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(value))

        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_file))

        return str(output_file)

    def fill_template(self, template_path: str, data: dict[str, str]) -> str:
        output_dir = Path("storage/outputs")
        output_dir.mkdir(parents=True, exist_ok=True)

        template_name = Path(template_path).name
        output_path = output_dir / f"{uuid4().hex}_filled_{template_name}"

        normalized_data = {k: str(v) for k, v in (data or {}).items()}
        return self.replace_placeholders(
            template_path=template_path,
            output_path=str(output_path),
            data=normalized_data,
        )
