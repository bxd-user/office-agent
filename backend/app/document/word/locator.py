"""
Word Document Locator module.
"""
from __future__ import annotations

import re
from typing import Any

from docx import Document


_PLACEHOLDER_PATTERNS = [
    r"\{\{(.*?)\}\}",
    r"【(.*?)】",
    r"<<(.*?)>>",
]


class WordLocator:
    """
    在 Word 文档中定位内容，支持：
    - 关键字搜索（段落 + 表格单元格）
    - 占位符定位
    - 表头行定位
    """

    def locate(
        self,
        file_path: str,
        keyword: str | None = None,
        placeholder: str | None = None,
        table_header: list[str] | None = None,
        case_sensitive: bool = False,
        field_name: str | None = None,
        table_cell: dict[str, int] | None = None,
        paragraph_index: int | None = None,
        nearby_window: int = 1,
        near_keyword: str | None = None,
        enable_fuzzy: bool = True,
        **kwargs,
    ) -> dict:
        doc = Document(file_path)
        attempts: list[dict[str, Any]] = []

        def run_strategy(name: str, fn) -> list[dict]:
            try:
                matches = fn()
                attempts.append({"strategy": name, "count": len(matches)})
                return matches
            except Exception as e:
                attempts.append({"strategy": name, "count": 0, "error": str(e)})
                return []

        # 1) 优先精确匹配策略
        strategy_chain: list[tuple[str, Any]] = []
        if placeholder is not None:
            strategy_chain.append(("exact_placeholder", lambda: self._locate_placeholder(doc, placeholder)))
        if table_cell is not None:
            strategy_chain.append(("table_cell", lambda: self._locate_table_cell(doc, table_cell)))
        if paragraph_index is not None or near_keyword is not None:
            strategy_chain.append(
                (
                    "paragraph_nearby",
                    lambda: self._locate_nearby_region(
                        doc,
                        paragraph_index=paragraph_index,
                        near_keyword=near_keyword,
                        window=max(0, int(nearby_window)),
                    ),
                )
            )
        if table_header is not None:
            strategy_chain.append(("table_header", lambda: self._locate_table_header(doc, table_header)))
        if keyword is not None:
            strategy_chain.append(("keyword", lambda: self._locate_keyword(doc, keyword, case_sensitive)))
        if field_name is not None and enable_fuzzy:
            strategy_chain.append(("fuzzy_field", lambda: self._locate_fuzzy_field(doc, field_name)))

        # 2) 若调用方没给参数，走默认扫描策略
        if not strategy_chain:
            strategy_chain = [
                ("all_placeholders", lambda: self._locate_all_placeholders(doc)),
            ]

        # 3) 失败降级：依次执行，找到第一个非空命中即返回
        for name, fn in strategy_chain:
            matches = run_strategy(name, fn)
            if matches:
                return {
                    "matches": matches,
                    "count": len(matches),
                    "strategy": name,
                    "attempts": attempts,
                }

        # 4) 最后兜底
        fallback = run_strategy("fallback_all_placeholders", lambda: self._locate_all_placeholders(doc))
        return {
            "matches": fallback,
            "count": len(fallback),
            "strategy": "fallback_all_placeholders",
            "attempts": attempts,
        }

    # ------------------------------------------------------------------
    # Locate strategies
    # ------------------------------------------------------------------

    def _locate_keyword(self, doc, keyword: str, case_sensitive: bool) -> list[dict]:
        matches: list[dict] = []

        def text_matches(text: str) -> bool:
            if case_sensitive:
                return keyword in text
            return keyword.lower() in text.lower()

        for i, p in enumerate(doc.paragraphs):
            if text_matches(p.text):
                matches.append({
                    "type": "paragraph",
                    "index": i,
                    "text": p.text,
                    "style": getattr(p.style, "name", None),
                })

        for ti, table in enumerate(doc.tables):
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    if text_matches(cell.text):
                        matches.append({
                            "type": "table_cell",
                            "table_index": ti,
                            "row_index": ri,
                            "col_index": ci,
                            "text": cell.text,
                        })

        return matches

    def _locate_fuzzy_field(self, doc, field_name: str) -> list[dict]:
        matches: list[dict] = []
        aliases = self._expand_field_aliases(field_name)
        aliases_norm = {self._normalize_text(x) for x in aliases}

        for i, p in enumerate(doc.paragraphs):
            text = p.text or ""
            norm = self._normalize_text(text)
            if any(alias in norm for alias in aliases_norm):
                score = max(self._fuzzy_score(norm, alias) for alias in aliases_norm)
                matches.append(
                    {
                        "type": "paragraph",
                        "index": i,
                        "text": text,
                        "style": getattr(p.style, "name", None),
                        "matched_field": field_name,
                        "match_mode": "fuzzy",
                        "score": round(score, 4),
                    }
                )

        for ti, table in enumerate(doc.tables):
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    text = cell.text or ""
                    norm = self._normalize_text(text)
                    if any(alias in norm for alias in aliases_norm):
                        score = max(self._fuzzy_score(norm, alias) for alias in aliases_norm)
                        matches.append(
                            {
                                "type": "table_cell",
                                "table_index": ti,
                                "row_index": ri,
                                "col_index": ci,
                                "text": text,
                                "matched_field": field_name,
                                "match_mode": "fuzzy",
                                "score": round(score, 4),
                            }
                        )

        matches.sort(key=lambda x: float(x.get("score", 0.0)), reverse=True)
        return matches

    def _locate_placeholder(self, doc, placeholder: str) -> list[dict]:
        """精确查找某个占位符名称（不含括号）出现的位置。"""
        matches: list[dict] = []
        targets = [
            f"{{{{{placeholder}}}}}",
            f"【{placeholder}】",
            f"<<{placeholder}>>",
        ]

        for i, p in enumerate(doc.paragraphs):
            for t in targets:
                if t in p.text:
                    matches.append({
                        "type": "paragraph",
                        "index": i,
                        "text": p.text,
                        "matched_token": t,
                    })
                    break

        for ti, table in enumerate(doc.tables):
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    for t in targets:
                        if t in cell.text:
                            matches.append({
                                "type": "table_cell",
                                "table_index": ti,
                                "row_index": ri,
                                "col_index": ci,
                                "text": cell.text,
                                "matched_token": t,
                            })
                            break

        return matches

    def _locate_table_header(self, doc, headers: list[str]) -> list[dict]:
        """找到包含指定表头的表格及其行号。"""
        matches: list[dict] = []
        normalized = [h.strip() for h in headers]

        for ti, table in enumerate(doc.tables):
            for ri, row in enumerate(table.rows):
                row_texts = [cell.text.strip() for cell in row.cells]
                found = [h for h in normalized if h in row_texts]
                if found:
                    matches.append({
                        "type": "table_header_row",
                        "table_index": ti,
                        "row_index": ri,
                        "row_texts": row_texts,
                        "matched_headers": found,
                        "match_score": len(found),
                    })

        matches.sort(key=lambda x: x["match_score"], reverse=True)
        return matches

    def _locate_table_cell(self, doc, table_cell: dict[str, int]) -> list[dict]:
        table_index = int(table_cell.get("table_index", -1))
        row_index = int(table_cell.get("row_index", -1))
        col_index = int(table_cell.get("col_index", -1))
        if table_index < 0 or row_index < 0 or col_index < 0:
            return []
        if table_index >= len(doc.tables):
            return []

        table = doc.tables[table_index]
        if row_index >= len(table.rows):
            return []
        row = table.rows[row_index]
        if col_index >= len(row.cells):
            return []

        cell = row.cells[col_index]
        return [
            {
                "type": "table_cell",
                "table_index": table_index,
                "row_index": row_index,
                "col_index": col_index,
                "text": cell.text,
                "match_mode": "exact_cell",
            }
        ]

    def _locate_nearby_region(
        self,
        doc,
        paragraph_index: int | None,
        near_keyword: str | None,
        window: int,
    ) -> list[dict]:
        base_indices: list[int] = []

        if paragraph_index is not None and 0 <= paragraph_index < len(doc.paragraphs):
            base_indices.append(paragraph_index)

        if near_keyword:
            for idx, p in enumerate(doc.paragraphs):
                if near_keyword in (p.text or ""):
                    base_indices.append(idx)

        if not base_indices:
            return []

        matches: list[dict] = []
        seen: set[int] = set()
        for base in base_indices:
            start = max(0, base - window)
            end = min(len(doc.paragraphs) - 1, base + window)
            for idx in range(start, end + 1):
                if idx in seen:
                    continue
                seen.add(idx)
                paragraph = doc.paragraphs[idx]
                matches.append(
                    {
                        "type": "paragraph_nearby",
                        "index": idx,
                        "base_index": base,
                        "distance": abs(idx - base),
                        "text": paragraph.text,
                        "style": getattr(paragraph.style, "name", None),
                    }
                )

        matches.sort(key=lambda x: (int(x.get("distance", 0)), int(x.get("index", 0))))
        return matches

    def _locate_all_placeholders(self, doc) -> list[dict]:
        """扫描文档内所有占位符并返回位置。"""
        matches: list[dict] = []

        def scan(text: str, location: dict):
            for pat in _PLACEHOLDER_PATTERNS:
                for m in re.finditer(pat, text):
                    matches.append({
                        "placeholder": m.group(0),
                        "key": m.group(1).strip(),
                        **location,
                    })

        for i, p in enumerate(doc.paragraphs):
            scan(p.text, {"type": "paragraph", "index": i})

        for ti, table in enumerate(doc.tables):
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    scan(cell.text, {
                        "type": "table_cell",
                        "table_index": ti,
                        "row_index": ri,
                        "col_index": ci,
                    })

        return matches

    def _expand_field_aliases(self, field_name: str) -> list[str]:
        base = str(field_name or "").strip()
        if not base:
            return []

        alias_map = {
            "name": ["姓名", "名称", "客户名称", "公司名称"],
            "phone": ["电话", "联系电话", "手机号", "手机"],
            "email": ["邮箱", "电子邮箱", "邮件"],
            "address": ["地址", "联系地址", "通讯地址"],
            "date": ["日期", "签署日期", "合同日期"],
            "amount": ["金额", "总金额", "费用", "价格"],
        }

        normalized = self._normalize_text(base)
        for key, aliases in alias_map.items():
            all_names = [key, *aliases]
            if any(self._normalize_text(name) == normalized for name in all_names):
                return all_names

        return [base]

    @staticmethod
    def _normalize_text(value: str) -> str:
        text = (value or "").strip().lower()
        text = text.replace(" ", "")
        text = text.replace("_", "")
        text = text.replace("-", "")
        text = text.replace(":", "")
        text = text.replace("：", "")
        return text

    @staticmethod
    def _fuzzy_score(text: str, query: str) -> float:
        if not text or not query:
            return 0.0
        if query in text:
            return min(1.0, len(query) / max(1, len(text)))

        # 简单字符重叠比分数
        text_chars = set(text)
        query_chars = set(query)
        overlap = len(text_chars & query_chars)
        return overlap / max(1, len(query_chars))
