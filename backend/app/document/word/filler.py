from __future__ import annotations

import os
import re
from typing import Any

from docx import Document


_PLACEHOLDER_TRIGGERS = ("{{", "【", "<<")


class WordFiller:
    def write_kv_pairs_to_template(
        self,
        file_path: str,
        mapping: dict[str, str],
        output_path: str,
        skip_empty_values: bool = True,
        overwrite_strategy: str = "replace",
        preserve_format: bool = True,
        table_cell_values: dict[str, Any] | None = None,
    ):
        doc = Document(file_path)
        replace_count = 0
        paragraph_replace_count = 0
        table_replace_count = 0
        skipped_keys: list[str] = []

        effective_mapping = self._build_effective_mapping(
            mapping=mapping,
            skip_empty_values=skip_empty_values,
            skipped_keys=skipped_keys,
        )

        for p in doc.paragraphs:
            changed = self._replace_in_paragraph(
                paragraph=p,
                mapping=effective_mapping,
                overwrite_strategy=overwrite_strategy,
                preserve_format=preserve_format,
            )
            replace_count += changed
            paragraph_replace_count += changed

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        changed = self._replace_in_paragraph(
                            paragraph=para,
                            mapping=effective_mapping,
                            overwrite_strategy=overwrite_strategy,
                            preserve_format=preserve_format,
                        )
                        replace_count += changed
                        table_replace_count += changed

        direct_cell_writes = self._fill_table_cells_directly(doc, table_cell_values or {}, overwrite_strategy)

        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)

        return {
            "output_path": output_path,
            "replace_count": replace_count,
            "paragraph_replace_count": paragraph_replace_count,
            "table_replace_count": table_replace_count,
            "direct_cell_write_count": direct_cell_writes,
            "used_mapping_keys": list(effective_mapping.keys()),
            "skipped_keys": skipped_keys,
            "overwrite_strategy": overwrite_strategy,
            "preserve_format": preserve_format,
        }

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    def _replace_in_paragraph(
        self,
        paragraph,
        mapping: dict[str, str],
        overwrite_strategy: str,
        preserve_format: bool,
    ) -> int:
        """
        替换段落中的占位符，正确处理跨 run 情况。

        策略：若段落文本含有占位符触发字符，先将所有 run 合并到首 run
        （保留首 run 的字体/样式），执行替换后返回替换次数。
        """
        full_text = paragraph.text
        if not any(trigger in full_text for trigger in _PLACEHOLDER_TRIGGERS):
            return 0

        # 优先 run 级替换，能最大限度保留模板样式
        run_replace_count = self._replace_in_runs(paragraph.runs, mapping)
        if run_replace_count > 0:
            return run_replace_count

        # 占位符跨 run 场景兜底：根据策略决定是否覆盖
        if preserve_format:
            # 尽量保留样式时，避免强制合并 run 导致样式丢失
            if overwrite_strategy == "keep_existing":
                return 0
            return 0

        # 不保留格式时，允许合并 run 以提高命中率
        self._normalize_runs(paragraph)

        count = 0
        for run in paragraph.runs:
            original = run.text
            new_text = original
            for key, value in mapping.items():
                for token in self._build_tokens(key):
                    if overwrite_strategy == "keep_existing" and not self._is_placeholder_token(token):
                        continue
                    new_text = new_text.replace(token, str(value))
            if new_text != original:
                run.text = new_text
                count += 1
        return count

    def _replace_in_runs(self, runs, mapping: dict[str, str]) -> int:
        count = 0
        for run in runs:
            original = run.text
            if not original:
                continue
            new_text = original
            for key, value in mapping.items():
                for token in self._build_tokens(key):
                    if token in new_text:
                        new_text = new_text.replace(token, str(value))
            if new_text != original:
                run.text = new_text
                count += 1
        return count

    def _fill_table_cells_directly(self, doc, table_cell_values: dict[str, Any], overwrite_strategy: str) -> int:
        if not table_cell_values:
            return 0

        write_count = 0
        for key, value in table_cell_values.items():
            parsed = self._parse_table_cell_key(str(key))
            if not parsed:
                continue
            table_index, row_index, col_index = parsed
            if table_index >= len(doc.tables):
                continue

            table = doc.tables[table_index]
            if row_index >= len(table.rows):
                continue
            row = table.rows[row_index]
            if col_index >= len(row.cells):
                continue

            cell = row.cells[col_index]
            existing = (cell.text or "").strip()
            if overwrite_strategy == "keep_existing" and existing:
                continue

            text_value = "" if value is None else str(value)
            self._write_cell_text_preserve_first_run(cell, text_value)
            write_count += 1

        return write_count

    @staticmethod
    def _parse_table_cell_key(key: str) -> tuple[int, int, int] | None:
        # 支持："0,1,2" / "0:1:2" / "t:0:r:1:c:2"
        compact = key.strip()
        m = re.match(r"^\s*(\d+)\s*[,\:]\s*(\d+)\s*[,\:]\s*(\d+)\s*$", compact)
        if m:
            return int(m.group(1)), int(m.group(2)), int(m.group(3))

        m2 = re.match(r"^t:(\d+):r:(\d+):c:(\d+)$", compact)
        if m2:
            return int(m2.group(1)), int(m2.group(2)), int(m2.group(3))
        return None

    @staticmethod
    def _build_effective_mapping(
        mapping: dict[str, Any],
        skip_empty_values: bool,
        skipped_keys: list[str],
    ) -> dict[str, str]:
        result: dict[str, str] = {}
        for key, value in (mapping or {}).items():
            if not isinstance(key, str):
                continue
            if value is None:
                if skip_empty_values:
                    skipped_keys.append(key)
                    continue
                result[key] = ""
                continue

            text = str(value)
            if skip_empty_values and text.strip() == "":
                skipped_keys.append(key)
                continue
            result[key] = text
        return result

    @staticmethod
    def _build_tokens(key: str) -> list[str]:
        safe = str(key).strip()
        if not safe:
            return []
        return [
            f"{{{{{safe}}}}}",
            f"【{safe}】",
            f"<<{safe}>>",
            f"{safe}:",
            f"{safe}：",
        ]

    @staticmethod
    def _is_placeholder_token(token: str) -> bool:
        return token.startswith("{{") or token.startswith("【") or token.startswith("<<")

    @staticmethod
    def _write_cell_text_preserve_first_run(cell, text: str) -> None:
        if not cell.paragraphs:
            cell.text = text
            return

        paragraph = cell.paragraphs[0]
        if not paragraph.runs:
            paragraph.text = text
            return

        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""

    @staticmethod
    def _normalize_runs(paragraph) -> None:
        """
        将段落所有 run 的文本合并到首 run，消除跨 run 的占位符分裂问题。
        只在段落含有占位符触发符时调用，避免不必要的格式损失。
        """
        runs = paragraph.runs
        if len(runs) <= 1:
            return

        full_text = "".join(r.text for r in runs)
        runs[0].text = full_text
        for run in runs[1:]:
            run.text = ""
