from __future__ import annotations

import re
from typing import Any

from docx import Document

from app.document.word.template_scanner import WordTemplateScanner


class WordAnalyzer:
    PLACEHOLDER_PATTERNS = [
        r"\{\{(.*?)\}\}",
        r"【(.*?)】",
        r"<<(.*?)>>",
    ]
    HEADING_STYLE_PATTERN = re.compile(r"(?:Heading|标题)\s*([1-9]\d*)", re.IGNORECASE)

    # 语义字段归一化映射：canonical -> aliases
    FIELD_ALIAS_MAP: dict[str, list[str]] = {
        "name": ["姓名", "客户名称", "公司名称", "单位名称", "名称"],
        "phone": ["电话", "联系电话", "手机号", "手机"],
        "email": ["邮箱", "电子邮箱", "邮件"],
        "address": ["地址", "联系地址", "通讯地址"],
        "date": ["日期", "签署日期", "合同日期", "填表日期"],
        "amount": ["金额", "总金额", "合同金额", "费用", "价格"],
        "id_number": ["身份证", "身份证号", "证件号", "统一社会信用代码"],
    }

    def __init__(self) -> None:
        self.template_scanner = WordTemplateScanner()

    def analyze_structure(self, file_path: str) -> dict[str, Any]:
        template = self.find_placeholders(file_path)
        hierarchy = self.analyze_heading_hierarchy(file_path)
        tables = self.analyze_table_structure(file_path)
        normalized = self.normalize_semantic_fields(template.get("fields", []))
        return {
            "template_fields": template,
            "heading_hierarchy": hierarchy,
            "table_analysis": tables,
            "semantic_normalization": normalized,
        }

    def find_placeholders(self, file_path: str):
        scan_result = self.template_scanner.scan(file_path)
        matches = []
        for item in scan_result.get("occurrences", []):
            location: dict[str, Any] = {"type": str(item.get("container") or "unknown")}
            for key in ("index", "table_index", "row_index", "col_index"):
                if key in item:
                    location[key] = item[key]
            matches.append(
                {
                    "placeholder": item.get("token") or item.get("field", ""),
                    "key": item.get("field", ""),
                    "location": location,
                }
            )

        normalized = self.normalize_semantic_fields(scan_result.get("fields", []))
        return {
            "placeholders": matches,
            "count": len(matches),
            "fields": scan_result.get("fields", []),
            "normalized_fields": normalized.get("normalized_fields", []),
            "field_mapping": normalized.get("field_mapping", {}),
        }

    def analyze_heading_hierarchy(self, file_path: str) -> dict[str, Any]:
        doc = Document(file_path)

        headings: list[dict[str, Any]] = []
        paragraphs: list[dict[str, Any]] = []
        current_path: dict[int, str] = {}

        for idx, paragraph in enumerate(doc.paragraphs):
            text = (paragraph.text or "").strip()
            if not text:
                continue

            style_name = getattr(paragraph.style, "name", "") or ""
            level = self._parse_heading_level(style_name, text)

            if level is not None:
                current_path[level] = text
                # 低优先级层级失效
                for key in list(current_path.keys()):
                    if key > level:
                        current_path.pop(key, None)
                heading_item = {
                    "index": idx,
                    "text": text,
                    "style": style_name,
                    "level": level,
                    "path": [current_path[k] for k in sorted(current_path.keys())],
                }
                headings.append(heading_item)

            paragraph_item = {
                "index": idx,
                "text": text,
                "style": style_name,
                "heading_level": level,
                "section_path": [current_path[k] for k in sorted(current_path.keys())],
            }
            paragraphs.append(paragraph_item)

        return {
            "headings": headings,
            "paragraphs": paragraphs,
            "heading_count": len(headings),
        }

    def analyze_table_structure(self, file_path: str) -> dict[str, Any]:
        doc = Document(file_path)
        analyses: list[dict[str, Any]] = []

        for table_index, table in enumerate(doc.tables):
            row_count = len(table.rows)
            col_count = max((len(row.cells) for row in table.rows), default=0)

            rows: list[list[str]] = []
            empty_cells: list[dict[str, int]] = []
            header_row_index = 0 if row_count > 0 else None

            for row_index, row in enumerate(table.rows):
                row_values: list[str] = []
                for col_index, cell in enumerate(row.cells):
                    text = (cell.text or "").strip()
                    row_values.append(text)
                    if not text:
                        empty_cells.append(
                            {
                                "row_index": row_index,
                                "col_index": col_index,
                            }
                        )
                rows.append(row_values)

            header_candidates = rows[header_row_index] if header_row_index is not None else []
            normalized_headers = [self._basic_normalize_text(x) for x in header_candidates if x]

            analyses.append(
                {
                    "table_index": table_index,
                    "row_count": row_count,
                    "col_count": col_count,
                    "header_row_index": header_row_index,
                    "header_candidates": header_candidates,
                    "normalized_headers": normalized_headers,
                    "empty_cells": empty_cells,
                    "rows": rows,
                }
            )

        return {
            "tables": analyses,
            "table_count": len(analyses),
        }

    def normalize_semantic_fields(self, fields: list[str]) -> dict[str, Any]:
        normalized_fields: list[dict[str, str]] = []
        field_mapping: dict[str, str] = {}
        for raw_field in fields:
            field = str(raw_field or "").strip()
            if not field:
                continue
            canonical = self._to_canonical_field(field)
            normalized_fields.append(
                {
                    "raw": field,
                    "canonical": canonical,
                }
            )
            field_mapping[field] = canonical

        return {
            "normalized_fields": normalized_fields,
            "field_mapping": field_mapping,
            "count": len(normalized_fields),
        }

    def find_paragraphs_by_keyword(self, file_path: str, keyword: str):
        doc = Document(file_path)
        items = []
        for i, p in enumerate(doc.paragraphs):
            if keyword in p.text:
                items.append({
                    "index": i,
                    "text": p.text,
                    "style": getattr(p.style, "name", None),
                })
        return {"matches": items, "count": len(items)}

    def find_table_by_header(self, file_path: str, headers: list[str]):
        doc = Document(file_path)
        candidates = []

        for ti, table in enumerate(doc.tables):
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            header_row = rows[0] if rows else []
            score = sum(1 for h in headers if h in header_row)

            if score > 0:
                candidates.append({
                    "table_index": ti,
                    "header_row": header_row,
                    "match_score": score,
                    "rows": rows,
                })

        candidates.sort(key=lambda x: x["match_score"], reverse=True)
        return {
            "candidates": candidates,
            "best_match": candidates[0] if candidates else None,
        }

    def _parse_heading_level(self, style_name: str, text: str) -> int | None:
        style_match = self.HEADING_STYLE_PATTERN.search(style_name or "")
        if style_match:
            return int(style_match.group(1))

        # 样式不规范时的弱规则回退：中文“一、/二、”、数字“1.”
        stripped = (text or "").strip()
        if re.match(r"^[一二三四五六七八九十]+、", stripped):
            return 1
        if re.match(r"^\d+[\.、]", stripped):
            return 2
        return None

    def _to_canonical_field(self, field: str) -> str:
        normalized_input = self._basic_normalize_text(field)

        for canonical, aliases in self.FIELD_ALIAS_MAP.items():
            if normalized_input == self._basic_normalize_text(canonical):
                return canonical
            for alias in aliases:
                if normalized_input == self._basic_normalize_text(alias):
                    return canonical

        return normalized_input

    @staticmethod
    def _basic_normalize_text(value: str) -> str:
        text = (value or "").strip().lower()
        text = text.replace(" ", "")
        text = text.replace("_", "")
        text = text.replace("-", "")
        return text