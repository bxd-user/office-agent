from __future__ import annotations

import re
from typing import Any

from docx import Document

from app.document.word.models import ValidationIssue, ValidationReport


class WordValidator:
    PLACEHOLDER_PATTERNS = [
        re.compile(r"\{\{\s*([^{}\n]+?)\s*\}\}"),
        re.compile(r"【\s*([^【】\n]+?)\s*】"),
        re.compile(r"<<\s*([^<>\n]+?)\s*>>"),
    ]

    def validate_replacements(
        self,
        file_path: str,
        source_template_path: str | None = None,
        expected_fields: list[str] | None = None,
        filled_values: dict[str, Any] | None = None,
        fill_targets: list[dict[str, Any]] | None = None,
    ):
        doc = Document(file_path)
        remaining = self._scan_remaining_placeholders(doc)
        issues: list[ValidationIssue] = []

        for item in remaining:
            issues.append(
                ValidationIssue(
                    issue_type="unfilled_placeholder",
                    severity="high",
                    message=f"存在未填占位符: {item.get('key') or item.get('text', '')}",
                    location_type=item.get("type"),
                    paragraph_index=item.get("index"),
                    table_index=item.get("table_index"),
                    row_index=item.get("row_index"),
                    col_index=item.get("col_index"),
                    details={"text": item.get("text")},
                )
            )

        empty_filled_fields = self._collect_empty_filled_fields(filled_values or {})
        for field in empty_filled_fields:
            issues.append(
                ValidationIssue(
                    issue_type="empty_filled_value",
                    severity="medium",
                    message=f"字段填充值为空: {field}",
                    details={"field": field},
                )
            )

        wrong_region_fills = self._detect_wrong_region_fills(doc, fill_targets or [])
        for item in wrong_region_fills:
            issues.append(
                ValidationIssue(
                    issue_type="wrong_region",
                    severity="high",
                    message=item.get("message", "疑似填错区域"),
                    location_type="table_cell",
                    table_index=item.get("table_index"),
                    row_index=item.get("row_index"),
                    col_index=item.get("col_index"),
                    details=item,
                )
            )

        structure_changes = self._detect_structure_change(
            source_template_path=source_template_path,
            result_doc=doc,
        )
        if structure_changes:
            issues.append(
                ValidationIssue(
                    issue_type="structure_change",
                    severity="high",
                    message="文档结构与模板差异明显，疑似写入破坏结构",
                    details=structure_changes,
                )
            )

        missing_expected_fields = self._collect_missing_expected_fields(
            expected_fields=expected_fields or [],
            remaining=remaining,
            filled_values=filled_values or {},
        )
        for field in missing_expected_fields:
            issues.append(
                ValidationIssue(
                    issue_type="missing_expected_field",
                    severity="medium",
                    message=f"存在明显缺漏字段: {field}",
                    details={"field": field},
                )
            )

        passed = len(issues) == 0
        report = ValidationReport(
            passed=passed,
            issues=issues,
            summary={
                "issue_count": len(issues),
                "remaining_placeholder_count": len(remaining),
                "empty_filled_count": len(empty_filled_fields),
                "wrong_region_count": len(wrong_region_fills),
                "missing_expected_count": len(missing_expected_fields),
                "has_structure_change": bool(structure_changes),
            },
            remaining_placeholders=remaining,
            empty_filled_fields=empty_filled_fields,
            wrong_region_fills=wrong_region_fills,
            missing_expected_fields=missing_expected_fields,
        )

        payload = self._model_to_dict(report)
        payload["count"] = len(remaining)
        payload["validation_passed"] = passed
        return payload

    def _scan_remaining_placeholders(self, doc) -> list[dict[str, Any]]:
        remaining: list[dict[str, Any]] = []

        def scan(text: str, location: dict[str, Any]) -> None:
            for pattern in self.PLACEHOLDER_PATTERNS:
                for m in pattern.finditer(text or ""):
                    remaining.append(
                        {
                            **location,
                            "text": text,
                            "token": m.group(0),
                            "key": str(m.group(1) or "").strip(),
                        }
                    )

        for i, p in enumerate(doc.paragraphs):
            scan(p.text, {"type": "paragraph", "index": i})

        for ti, table in enumerate(doc.tables):
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    scan(
                        cell.text,
                        {
                            "type": "table_cell",
                            "table_index": ti,
                            "row_index": ri,
                            "col_index": ci,
                        },
                    )

        return remaining

    @staticmethod
    def _collect_empty_filled_fields(filled_values: dict[str, Any]) -> list[str]:
        empty_fields: list[str] = []
        for key, value in filled_values.items():
            if value is None:
                empty_fields.append(str(key))
                continue
            if isinstance(value, str) and not value.strip():
                empty_fields.append(str(key))
        return empty_fields

    @staticmethod
    def _detect_wrong_region_fills(doc, fill_targets: list[dict[str, Any]]) -> list[dict[str, Any]]:
        issues: list[dict[str, Any]] = []
        for target in fill_targets:
            table_cell = target.get("table_cell")
            value = target.get("value")
            if not isinstance(table_cell, dict):
                continue
            try:
                ti = int(table_cell.get("table_index"))
                ri = int(table_cell.get("row_index"))
                ci = int(table_cell.get("col_index"))
            except Exception:
                continue

            if ti < 0 or ri < 0 or ci < 0 or ti >= len(doc.tables):
                issues.append(
                    {
                        "message": "填充目标单元格越界",
                        "table_index": ti,
                        "row_index": ri,
                        "col_index": ci,
                    }
                )
                continue

            table = doc.tables[ti]
            if ri >= len(table.rows) or ci >= len(table.rows[ri].cells):
                issues.append(
                    {
                        "message": "填充目标单元格越界",
                        "table_index": ti,
                        "row_index": ri,
                        "col_index": ci,
                    }
                )
                continue

            cell_text = (table.rows[ri].cells[ci].text or "").strip()
            expected = "" if value is None else str(value).strip()
            if expected and expected not in cell_text:
                issues.append(
                    {
                        "message": "目标区域未出现期望填充值，疑似填错区域",
                        "table_index": ti,
                        "row_index": ri,
                        "col_index": ci,
                        "expected": expected,
                        "actual": cell_text,
                    }
                )

        return issues

    @staticmethod
    def _detect_structure_change(source_template_path: str | None, result_doc) -> dict[str, Any] | None:
        if not source_template_path:
            return None

        try:
            src = Document(source_template_path)
        except Exception:
            return None

        src_summary = {
            "paragraph_count": len(src.paragraphs),
            "table_count": len(src.tables),
            "table_rows": sum(len(t.rows) for t in src.tables),
        }
        dst_summary = {
            "paragraph_count": len(result_doc.paragraphs),
            "table_count": len(result_doc.tables),
            "table_rows": sum(len(t.rows) for t in result_doc.tables),
        }

        differences = {}
        for key in src_summary:
            if src_summary[key] != dst_summary[key]:
                differences[key] = {"source": src_summary[key], "result": dst_summary[key]}

        return differences or None

    @staticmethod
    def _collect_missing_expected_fields(
        expected_fields: list[str],
        remaining: list[dict[str, Any]],
        filled_values: dict[str, Any],
    ) -> list[str]:
        expected = {str(x).strip() for x in expected_fields if str(x).strip()}
        if not expected:
            return []

        remaining_keys = {str(item.get("key") or "").strip() for item in remaining}
        filled_keys = {str(k).strip() for k in filled_values.keys()}

        missing: list[str] = []
        for field in sorted(expected):
            if field in remaining_keys:
                missing.append(field)
                continue
            if field not in filled_keys:
                missing.append(field)
                continue
            value = filled_values.get(field)
            if value is None or (isinstance(value, str) and not value.strip()):
                missing.append(field)
        return missing

    @staticmethod
    def _model_to_dict(model: Any) -> dict[str, Any]:
        if hasattr(model, "model_dump"):
            return model.model_dump()
        if hasattr(model, "dict"):
            return model.dict()
        return dict(model)