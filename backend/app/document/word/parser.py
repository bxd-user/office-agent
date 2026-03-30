from __future__ import annotations

import re
from typing import Any

from docx import Document


PLACEHOLDER_PATTERN = re.compile(r"\{\{(.*?)\}\}")
BRACKET_LABEL_PATTERN = re.compile(r"【(.*?)】|<<(.*?)>>")
INLINE_LABEL_PATTERN = re.compile(r"([\u4e00-\u9fa5A-Za-z0-9_\-]{1,30})\s*[:：]")


def extract_placeholders(text: str) -> list[str]:
    if not text:
        return []
    return [x.strip() for x in PLACEHOLDER_PATTERN.findall(text) if x.strip()]


class WordParser:
    def read_text(self, file_path: str) -> str:
        structure = self.extract_structure(file_path)
        parts: list[str] = []
        for paragraph in structure.get("paragraphs", []):
            text = str(paragraph.get("text") or "").strip()
            if text:
                parts.append(text)
        return "\n".join(parts)

    def read_tables(self, file_path: str):
        structure = self.extract_structure(file_path)
        tables = structure.get("tables", [])
        # 向后兼容：保留 table_index + rows，并附加 cells/tag_candidates/position_index
        return tables

    def extract_structure(self, file_path: str):
        doc = Document(file_path)

        paragraphs: list[dict[str, Any]] = []
        tables: list[dict[str, Any]] = []
        runs: list[dict[str, Any]] = []
        tag_candidates: list[dict[str, Any]] = []
        position_index: dict[str, dict[str, Any]] = {}

        for i, p in enumerate(doc.paragraphs):
            paragraph_key = f"p:{i}"
            paragraph_runs: list[dict[str, Any]] = []
            for run_index, run in enumerate(p.runs):
                run_key = f"p:{i}:r:{run_index}"
                run_item = {
                    "run_index": run_index,
                    "text": run.text,
                    "bold": bool(run.bold) if run.bold is not None else False,
                    "italic": bool(run.italic) if run.italic is not None else False,
                    "underline": bool(run.underline) if run.underline is not None else False,
                    "position_key": run_key,
                }
                paragraph_runs.append(run_item)
                runs.append(
                    {
                        "owner_type": "paragraph",
                        "owner_index": i,
                        **run_item,
                    }
                )
                position_index[run_key] = {
                    "type": "run",
                    "owner_type": "paragraph",
                    "paragraph_index": i,
                    "run_index": run_index,
                }

            paragraph_tags = self._collect_tag_candidates(p.text)
            for tag in paragraph_tags:
                tag_candidates.append(
                    {
                        "tag": tag,
                        "source_type": "paragraph",
                        "paragraph_index": i,
                        "position_key": paragraph_key,
                    }
                )

            paragraphs.append({
                "index": i,
                "text": p.text,
                "style": getattr(p.style, "name", None),
                "runs": paragraph_runs,
                "tag_candidates": paragraph_tags,
                "position_key": paragraph_key,
            })
            position_index[paragraph_key] = {
                "type": "paragraph",
                "paragraph_index": i,
            }

        for i, table in enumerate(doc.tables):
            table_key = f"t:{i}"
            rows: list[list[str]] = []
            cells: list[dict[str, Any]] = []

            for row_index, row in enumerate(table.rows):
                row_values: list[str] = []
                for col_index, cell in enumerate(row.cells):
                    cell_text = cell.text
                    row_values.append(cell_text)
                    cell_key = f"t:{i}:r:{row_index}:c:{col_index}"
                    cell_tags = self._collect_tag_candidates(cell_text)
                    cells.append(
                        {
                            "row_index": row_index,
                            "col_index": col_index,
                            "text": cell_text,
                            "tag_candidates": cell_tags,
                            "position_key": cell_key,
                        }
                    )
                    for tag in cell_tags:
                        tag_candidates.append(
                            {
                                "tag": tag,
                                "source_type": "table_cell",
                                "table_index": i,
                                "row_index": row_index,
                                "col_index": col_index,
                                "position_key": cell_key,
                            }
                        )
                    position_index[cell_key] = {
                        "type": "table_cell",
                        "table_index": i,
                        "row_index": row_index,
                        "col_index": col_index,
                    }

                    for paragraph in cell.paragraphs:
                        for run_index, run in enumerate(paragraph.runs):
                            run_key = f"{cell_key}:r:{run_index}"
                            run_item = {
                                "run_index": run_index,
                                "text": run.text,
                                "bold": bool(run.bold) if run.bold is not None else False,
                                "italic": bool(run.italic) if run.italic is not None else False,
                                "underline": bool(run.underline) if run.underline is not None else False,
                                "position_key": run_key,
                            }
                            runs.append(
                                {
                                    "owner_type": "table_cell",
                                    "table_index": i,
                                    "row_index": row_index,
                                    "col_index": col_index,
                                    **run_item,
                                }
                            )
                            position_index[run_key] = {
                                "type": "run",
                                "owner_type": "table_cell",
                                "table_index": i,
                                "row_index": row_index,
                                "col_index": col_index,
                                "run_index": run_index,
                            }

                rows.append(row_values)

            tables.append({
                "table_index": i,
                "rows": rows,
                "cells": cells,
                "position_key": table_key,
            })
            position_index[table_key] = {
                "type": "table",
                "table_index": i,
            }

        unique_tag_candidates = self._dedupe_tags(tag_candidates)
        return {
            "paragraphs": paragraphs,
            "tables": tables,
            "runs": runs,
            "tag_candidates": unique_tag_candidates,
            "position_index": position_index,
        }

    @staticmethod
    def _collect_tag_candidates(text: str) -> list[str]:
        if not text:
            return []

        candidates: list[str] = []
        candidates.extend(extract_placeholders(text))

        for match in BRACKET_LABEL_PATTERN.findall(text):
            for item in match:
                clean = item.strip()
                if clean:
                    candidates.append(clean)

        for item in INLINE_LABEL_PATTERN.findall(text):
            clean = item.strip()
            if clean:
                candidates.append(clean)

        # 去重并保持顺序
        seen: set[str] = set()
        result: list[str] = []
        for candidate in candidates:
            if candidate not in seen:
                seen.add(candidate)
                result.append(candidate)
        return result

    @staticmethod
    def _dedupe_tags(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
        seen: set[tuple[str, str, str]] = set()
        result: list[dict[str, Any]] = []
        for item in items:
            key = (
                str(item.get("tag") or ""),
                str(item.get("source_type") or ""),
                str(item.get("position_key") or ""),
            )
            if key in seen:
                continue
            seen.add(key)
            result.append(item)
        return result