"""
Word Document Comparator module.
"""
from __future__ import annotations

import difflib
import re
from typing import Any

from docx import Document

from app.document.word.models import StructureSummaryDiff, TableDiff, WordComparisonReport
from app.document.word.template_scanner import WordTemplateScanner


class WordComparator:
    """
    对比两个 Word 文档，返回段落级别的差异。

    结果包含：
    - added   : 仅在右侧文档出现的段落
    - removed : 仅在左侧文档出现的段落
    - changed : 行号位置发生变化的段落对
    - summary : 统计数字
    """

    def compare(
        self,
        left_file_path: str,
        right_file_path: str,
        include_tables: bool = True,
        **kwargs,
    ) -> dict:
        left_doc = Document(left_file_path)
        right_doc = Document(right_file_path)
        scanner = WordTemplateScanner()

        left_paras = self._extract_text_blocks(left_doc, include_tables)
        right_paras = self._extract_text_blocks(right_doc, include_tables)

        text_diff = self._diff(left_paras, right_paras)
        field_diff = self._compare_fields(
            left_scan=scanner.scan(left_file_path),
            right_scan=scanner.scan(right_file_path),
        )
        table_diff = self._compare_tables(left_doc=left_doc, right_doc=right_doc)
        structure_summary_diff = self._compare_structure_summary(left_doc=left_doc, right_doc=right_doc)

        report = WordComparisonReport(
            left_file=left_file_path,
            right_file=right_file_path,
            text_diff={
                "added": text_diff["added"],
                "removed": text_diff["removed"],
                "changed": text_diff["changed"],
                "unchanged_count": text_diff["unchanged_count"],
            },
            field_diff=field_diff,
            table_diff=table_diff,
            structure_summary_diff=structure_summary_diff,
            summary={
                "text_changed_count": len(text_diff["changed"]),
                "text_added_count": len(text_diff["added"]),
                "text_removed_count": len(text_diff["removed"]),
                "field_added_count": len(field_diff.get("added_fields", [])),
                "field_removed_count": len(field_diff.get("removed_fields", [])),
                "field_common_count": len(field_diff.get("common_fields", [])),
                "table_changed_count": len([x for x in table_diff if x.changed_cells]),
                "structure_changed": bool(structure_summary_diff.differences),
            },
        )

        payload = self._model_to_dict(report)

        return {
            "left_file": left_file_path,
            "right_file": right_file_path,
            "text_diff": payload.get("text_diff", {}),
            "field_diff": payload.get("field_diff", {}),
            "table_diff": payload.get("table_diff", []),
            "structure_summary_diff": payload.get("structure_summary_diff", {}),
            "summary": payload.get("summary", {}),
        }

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    def _extract_text_blocks(self, doc, include_tables: bool) -> list[str]:
        blocks: list[str] = []

        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                blocks.append(text)

        if include_tables:
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip(" |"):
                        blocks.append(f"[TABLE] {row_text}")

        return blocks

    def _diff(self, left: list[str], right: list[str]) -> dict:
        matcher = difflib.SequenceMatcher(None, left, right, autojunk=False)

        added: list[dict] = []
        removed: list[dict] = []
        changed: list[dict] = []
        unchanged_count = 0

        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == "equal":
                unchanged_count += i2 - i1
            elif tag == "insert":
                for j in range(j1, j2):
                    added.append({"right_index": j, "text": right[j]})
            elif tag == "delete":
                for i in range(i1, i2):
                    removed.append({"left_index": i, "text": left[i]})
            elif tag == "replace":
                left_chunk = left[i1:i2]
                right_chunk = right[j1:j2]
                # 逐行配对，多余的归入 added/removed
                for k in range(max(len(left_chunk), len(right_chunk))):
                    l_text = left_chunk[k] if k < len(left_chunk) else None
                    r_text = right_chunk[k] if k < len(right_chunk) else None
                    if l_text and r_text:
                        changed.append({
                            "left_index": i1 + k,
                            "right_index": j1 + k,
                            "left_text": l_text,
                            "right_text": r_text,
                            "similarity": round(
                                difflib.SequenceMatcher(None, l_text, r_text).ratio(), 2
                            ),
                        })
                    elif l_text:
                        removed.append({"left_index": i1 + k, "text": l_text})
                    elif r_text:
                        added.append({"right_index": j1 + k, "text": r_text})

        return {
            "added": added,
            "removed": removed,
            "changed": changed,
            "unchanged_count": unchanged_count,
        }

    def _compare_fields(self, left_scan: dict[str, Any], right_scan: dict[str, Any]) -> dict[str, Any]:
        left_fields = set(left_scan.get("fields", []))
        right_fields = set(right_scan.get("fields", []))

        return {
            "left_fields": sorted(left_fields),
            "right_fields": sorted(right_fields),
            "added_fields": sorted(right_fields - left_fields),
            "removed_fields": sorted(left_fields - right_fields),
            "common_fields": sorted(left_fields & right_fields),
        }

    def _compare_tables(self, left_doc, right_doc) -> list[TableDiff]:
        left_tables = self._extract_tables(left_doc)
        right_tables = self._extract_tables(right_doc)

        max_count = max(len(left_tables), len(right_tables))
        result: list[TableDiff] = []
        for i in range(max_count):
            left_rows = left_tables[i] if i < len(left_tables) else []
            right_rows = right_tables[i] if i < len(right_tables) else []
            changed_cells: list[dict[str, Any]] = []

            row_max = max(len(left_rows), len(right_rows))
            for r in range(row_max):
                left_row = left_rows[r] if r < len(left_rows) else []
                right_row = right_rows[r] if r < len(right_rows) else []
                col_max = max(len(left_row), len(right_row))
                for c in range(col_max):
                    left_cell = left_row[c] if c < len(left_row) else ""
                    right_cell = right_row[c] if c < len(right_row) else ""
                    if left_cell != right_cell:
                        changed_cells.append(
                            {
                                "row_index": r,
                                "col_index": c,
                                "left": left_cell,
                                "right": right_cell,
                            }
                        )

            result.append(
                TableDiff(
                    table_index=i,
                    changed_cells=changed_cells,
                    row_count_left=len(left_rows),
                    row_count_right=len(right_rows),
                )
            )
        return result

    def _compare_structure_summary(self, left_doc, right_doc) -> StructureSummaryDiff:
        left_summary = self._build_structure_summary(left_doc)
        right_summary = self._build_structure_summary(right_doc)

        differences: dict[str, Any] = {}
        for key in sorted(set(left_summary.keys()) | set(right_summary.keys())):
            lv = left_summary.get(key)
            rv = right_summary.get(key)
            if lv != rv:
                differences[key] = {"left": lv, "right": rv}

        return StructureSummaryDiff(
            left_summary=left_summary,
            right_summary=right_summary,
            differences=differences,
        )

    def _extract_tables(self, doc) -> list[list[list[str]]]:
        tables: list[list[list[str]]] = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([(cell.text or "").strip() for cell in row.cells])
            tables.append(rows)
        return tables

    def _build_structure_summary(self, doc) -> dict[str, Any]:
        paragraph_count = len(doc.paragraphs)
        non_empty_paragraphs = sum(1 for p in doc.paragraphs if (p.text or "").strip())
        heading_count = 0
        for p in doc.paragraphs:
            style_name = str(getattr(getattr(p, "style", None), "name", "") or "")
            text = str(p.text or "").strip()
            if re.search(r"(?:Heading|标题)\s*\d+", style_name, re.IGNORECASE) or re.match(r"^[一二三四五六七八九十]+、", text):
                heading_count += 1

        table_count = len(doc.tables)
        table_rows = sum(len(t.rows) for t in doc.tables)
        table_cells = sum(len(r.cells) for t in doc.tables for r in t.rows)

        return {
            "paragraph_count": paragraph_count,
            "non_empty_paragraph_count": non_empty_paragraphs,
            "heading_count": heading_count,
            "table_count": table_count,
            "table_row_count": table_rows,
            "table_cell_count": table_cells,
        }

    @staticmethod
    def _model_to_dict(model: Any) -> dict[str, Any]:
        if hasattr(model, "model_dump"):
            return model.model_dump()
        if hasattr(model, "dict"):
            return model.dict()
        return dict(model)
