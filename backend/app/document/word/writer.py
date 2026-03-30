from __future__ import annotations

from copy import deepcopy
import os
from typing import Any

from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from app.document.word.models import WriteOutputMeta


class WordWriter:
    def _find_reference_paragraph_and_run(self, doc: Document) -> tuple[Paragraph | None, Run | None]:
        for paragraph in reversed(doc.paragraphs):
            if not paragraph.text or not paragraph.text.strip():
                continue
            for run in reversed(paragraph.runs):
                if run.text and run.text.strip():
                    return paragraph, run
            return paragraph, None

        for table in reversed(doc.tables):
            for row in reversed(table.rows):
                for cell in reversed(row.cells):
                    for paragraph in reversed(cell.paragraphs):
                        if not paragraph.text or not paragraph.text.strip():
                            continue
                        for run in reversed(paragraph.runs):
                            if run.text and run.text.strip():
                                return paragraph, run
                        return paragraph, None

        return None, None

    def _copy_run_style(self, src: Run, dst: Run) -> None:
        dst.bold = src.bold
        dst.italic = src.italic
        dst.underline = src.underline
        dst.style = src.style

        dst.font.name = src.font.name
        dst.font.size = src.font.size
        dst.font.bold = src.font.bold
        dst.font.italic = src.font.italic
        dst.font.underline = src.font.underline
        dst.font.highlight_color = src.font.highlight_color
        dst.font.strike = src.font.strike
        dst.font.subscript = src.font.subscript
        dst.font.superscript = src.font.superscript

        if src.font.color is not None:
            dst.font.color.rgb = src.font.color.rgb

    def replace_text(
        self,
        file_path: str,
        replacements: dict[str, str],
        output_path: str | None = None,
        output_dir: str | None = None,
        filename_prefix: str | None = None,
        filename_suffix: str = "_written",
        avoid_overwrite: bool = True,
    ):
        doc = Document(file_path)
        replaced_count = 0

        for p in doc.paragraphs:
            for old, new in replacements.items():
                if old in p.text:
                    before = p.text
                    for run in p.runs:
                        run.text = run.text.replace(old, new)
                    if p.text != before:
                        replaced_count += 1

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for old, new in replacements.items():
                        if old in cell.text:
                            before = cell.text
                            cell.text = cell.text.replace(old, new)
                            if cell.text != before:
                                replaced_count += 1

        resolved_output_path, avoided = self._resolve_output_path(
            source_file_path=file_path,
            output_path=output_path,
            output_dir=output_dir,
            filename_prefix=filename_prefix,
            filename_suffix=filename_suffix,
            avoid_overwrite=avoid_overwrite,
        )

        os.makedirs(os.path.dirname(resolved_output_path), exist_ok=True)
        doc.save(resolved_output_path)

        meta = WriteOutputMeta(
            source_file=file_path,
            output_file=resolved_output_path,
            output_dir=os.path.dirname(resolved_output_path),
            output_filename=os.path.basename(resolved_output_path),
            operation="replace_text",
            avoided_overwrite=avoided,
            replaced_count=replaced_count,
            appended=False,
        )

        return {
            "output_path": resolved_output_path,
            "replacements_count": len(replacements),
            "effective_replaced_blocks": replaced_count,
            "meta": self._model_to_dict(meta),
        }

    def append_text(
        self,
        file_path: str,
        append_text: str,
        output_path: str | None = None,
        output_dir: str | None = None,
        filename_prefix: str | None = None,
        filename_suffix: str = "_written",
        avoid_overwrite: bool = True,
    ):
        doc = Document(file_path)

        text = (append_text or "").strip()
        if text:
            reference_paragraph, reference_run = self._find_reference_paragraph_and_run(doc)
            doc.add_paragraph("")
            new_paragraph = doc.add_paragraph()
            if reference_paragraph is not None:
                new_paragraph.style = reference_paragraph.style

            new_run = new_paragraph.add_run(text)
            if reference_run is not None:
                self._copy_run_style(reference_run, new_run)

        resolved_output_path, avoided = self._resolve_output_path(
            source_file_path=file_path,
            output_path=output_path,
            output_dir=output_dir,
            filename_prefix=filename_prefix,
            filename_suffix=filename_suffix,
            avoid_overwrite=avoid_overwrite,
        )

        os.makedirs(os.path.dirname(resolved_output_path), exist_ok=True)
        doc.save(resolved_output_path)

        meta = WriteOutputMeta(
            source_file=file_path,
            output_file=resolved_output_path,
            output_dir=os.path.dirname(resolved_output_path),
            output_filename=os.path.basename(resolved_output_path),
            operation="append_text",
            avoided_overwrite=avoided,
            replaced_count=0,
            appended=bool(text),
        )

        return {
            "output_path": resolved_output_path,
            "appended": bool(text),
            "meta": self._model_to_dict(meta),
        }

    def save_as_result(
        self,
        file_path: str,
        output_path: str | None = None,
        output_dir: str | None = None,
        filename_prefix: str | None = None,
        filename_suffix: str = "_result",
        avoid_overwrite: bool = True,
    ) -> dict[str, Any]:
        doc = Document(file_path)
        resolved_output_path, avoided = self._resolve_output_path(
            source_file_path=file_path,
            output_path=output_path,
            output_dir=output_dir,
            filename_prefix=filename_prefix,
            filename_suffix=filename_suffix,
            avoid_overwrite=avoid_overwrite,
        )

        os.makedirs(os.path.dirname(resolved_output_path), exist_ok=True)
        doc.save(resolved_output_path)

        meta = WriteOutputMeta(
            source_file=file_path,
            output_file=resolved_output_path,
            output_dir=os.path.dirname(resolved_output_path),
            output_filename=os.path.basename(resolved_output_path),
            operation="save_as",
            avoided_overwrite=avoided,
            replaced_count=0,
            appended=False,
        )
        return {
            "output_path": resolved_output_path,
            "saved": True,
            "meta": self._model_to_dict(meta),
        }

    def _resolve_output_path(
        self,
        source_file_path: str,
        output_path: str | None,
        output_dir: str | None,
        filename_prefix: str | None,
        filename_suffix: str,
        avoid_overwrite: bool,
    ) -> tuple[str, bool]:
        source_abs = os.path.abspath(source_file_path)
        source_dir = os.path.dirname(source_abs)
        source_name = os.path.basename(source_abs)
        source_stem, source_ext = os.path.splitext(source_name)

        target_dir = output_dir or os.path.join(source_dir, "outputs")
        if output_path:
            candidate = os.path.abspath(output_path)
        else:
            prefix = f"{filename_prefix}_" if filename_prefix else ""
            filename = f"{prefix}{source_stem}{filename_suffix}{source_ext}"
            candidate = os.path.abspath(os.path.join(target_dir, filename))

        avoided = False
        if avoid_overwrite and (os.path.abspath(candidate) == source_abs or os.path.exists(candidate)):
            candidate = self._next_available_path(candidate)
            avoided = True

        # 确保不是原文件
        if os.path.abspath(candidate) == source_abs:
            candidate = self._next_available_path(candidate)
            avoided = True

        return candidate, avoided

    @staticmethod
    def _next_available_path(path: str) -> str:
        base, ext = os.path.splitext(path)
        index = 1
        candidate = f"{base}_{index}{ext}"
        while os.path.exists(candidate):
            index += 1
            candidate = f"{base}_{index}{ext}"
        return candidate

    @staticmethod
    def _model_to_dict(model: Any) -> dict[str, Any]:
        if hasattr(model, "model_dump"):
            return model.model_dump()
        if hasattr(model, "dict"):
            return model.dict()
        return dict(model)


DocxWriter = WordWriter
