from __future__ import annotations

import re
from typing import Any

from docx import Document


class WordTemplateScanner:
    PATTERNS = [
        re.compile(r"\{\{\s*([^\{\}\n]+?)\s*\}\}"),
        re.compile(r"\{\s*([^\{\}\n]+?)\s*\}"),
        re.compile(r"<<\s*([^<>\n]+?)\s*>>"),
        re.compile(r"【\s*([^【】\n]+?)\s*】"),
    ]
    LABEL_VALUE_PATTERN = re.compile(
        r"([\u4e00-\u9fa5A-Za-z][\u4e00-\u9fa5A-Za-z0-9_\-（）()]{0,30})\s*[:：]\s*([^\n]{0,80})"
    )
    CHINESE_FIELD_PATTERN = re.compile(
        r"(姓名|客户名称|公司名称|电话|手机号|联系地址|地址|金额|总金额|日期|签署日期|编号|合同编号|身份证号|证件号)"
    )

    def scan(self, file_path: str) -> dict[str, Any]:
        doc = Document(file_path)

        hits: list[dict[str, Any]] = []
        seen: set[str] = set()
        empty_table_cells: list[dict[str, Any]] = []
        label_value_regions: list[dict[str, Any]] = []
        chinese_field_candidates: list[dict[str, Any]] = []

        for idx, para in enumerate(doc.paragraphs):
            text = para.text or ""
            for field, token in self._extract_fields_with_token(text):
                if field not in seen:
                    seen.add(field)
                hits.append(
                    {
                        "container": "paragraph",
                        "index": idx,
                        "field": field,
                        "token": token,
                        "text": text,
                    }
                )

            label_value_regions.extend(
                self._scan_label_value_regions(
                    text=text,
                    location={"container": "paragraph", "index": idx},
                )
            )
            chinese_field_candidates.extend(
                self._scan_chinese_field_candidates(
                    text=text,
                    location={"container": "paragraph", "index": idx},
                )
            )

        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    text = cell.text or ""
                    if not text.strip():
                        empty_table_cells.append(
                            {
                                "table_index": t_idx,
                                "row_index": r_idx,
                                "col_index": c_idx,
                                "container": "table_cell",
                            }
                        )

                    for field, token in self._extract_fields_with_token(text):
                        if field not in seen:
                            seen.add(field)
                        hits.append(
                            {
                                "container": "table_cell",
                                "table_index": t_idx,
                                "row_index": r_idx,
                                "col_index": c_idx,
                                "field": field,
                                "token": token,
                                "text": text,
                            }
                        )

                    label_value_regions.extend(
                        self._scan_label_value_regions(
                            text=text,
                            location={
                                "container": "table_cell",
                                "table_index": t_idx,
                                "row_index": r_idx,
                                "col_index": c_idx,
                            },
                        )
                    )
                    chinese_field_candidates.extend(
                        self._scan_chinese_field_candidates(
                            text=text,
                            location={
                                "container": "table_cell",
                                "table_index": t_idx,
                                "row_index": r_idx,
                                "col_index": c_idx,
                            },
                        )
                    )

                # 识别常见标签-值区域：同一行相邻两列
                label_value_regions.extend(
                    self._scan_row_label_value_pairs(table_index=t_idx, row_index=r_idx, row=row)
                )

        fields = sorted(seen)

        return {
            "fields": fields,
            "field_count": len(fields),
            "occurrences": hits,
            "empty_table_cells": empty_table_cells,
            "empty_table_cell_count": len(empty_table_cells),
            "label_value_regions": self._dedupe_regions(label_value_regions),
            "chinese_field_candidates": self._dedupe_regions(chinese_field_candidates),
        }

    def _extract_fields(self, text: str) -> list[str]:
        return [item[0] for item in self._extract_fields_with_token(text)]

    def _extract_fields_with_token(self, text: str) -> list[tuple[str, str]]:
        results: list[str] = []
        for pattern in self.PATTERNS:
            for match in pattern.finditer(text):
                value = str(match.group(1)).strip()
                if value:
                    results.append((value, match.group(0)))
        return results

    def _scan_label_value_regions(self, text: str, location: dict[str, Any]) -> list[dict[str, Any]]:
        results: list[dict[str, Any]] = []
        if not text:
            return results

        for match in self.LABEL_VALUE_PATTERN.finditer(text):
            label = str(match.group(1) or "").strip()
            value = str(match.group(2) or "").strip()
            if not label:
                continue
            item = {
                **location,
                "label": label,
                "value": value,
                "pattern": "inline_label_value",
                "raw": match.group(0),
            }
            results.append(item)
        return results

    def _scan_row_label_value_pairs(self, table_index: int, row_index: int, row) -> list[dict[str, Any]]:
        results: list[dict[str, Any]] = []
        cells = list(row.cells)
        for col_index in range(0, len(cells) - 1):
            label = (cells[col_index].text or "").strip()
            value = (cells[col_index + 1].text or "").strip()
            if not label:
                continue

            # 弱规则：左侧是字段标签，右侧为空或非标签化内容
            if self._looks_like_label(label):
                results.append(
                    {
                        "container": "table_row_pair",
                        "table_index": table_index,
                        "row_index": row_index,
                        "label_col_index": col_index,
                        "value_col_index": col_index + 1,
                        "label": label,
                        "value": value,
                        "pattern": "adjacent_cells",
                    }
                )
        return results

    def _scan_chinese_field_candidates(self, text: str, location: dict[str, Any]) -> list[dict[str, Any]]:
        results: list[dict[str, Any]] = []
        if not text:
            return results
        for match in self.CHINESE_FIELD_PATTERN.finditer(text):
            field = str(match.group(1) or "").strip()
            if not field:
                continue
            results.append(
                {
                    **location,
                    "field": field,
                    "pattern": "chinese_field_keyword",
                    "raw": match.group(0),
                }
            )
        return results

    @staticmethod
    def _looks_like_label(text: str) -> bool:
        value = (text or "").strip()
        if not value:
            return False
        if len(value) > 20:
            return False
        if any(token in value for token in ("{{", "}}", "<<", ">>", "【", "】")):
            return False
        return bool(re.search(r"[\u4e00-\u9fa5A-Za-z]", value))

    @staticmethod
    def _dedupe_regions(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
        seen: set[tuple[str, ...]] = set()
        result: list[dict[str, Any]] = []
        for item in items:
            signature = tuple(sorted((str(k), str(v)) for k, v in item.items()))
            if signature in seen:
                continue
            seen.add(signature)
            result.append(item)
        return result