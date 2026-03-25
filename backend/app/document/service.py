import shutil
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from fastapi import UploadFile

from app.agent.loop import WorkflowAgent
from app.document.router import DocumentRouter
from app.document.word.parser import PLACEHOLDER_PATTERN
from app.domain.models import InputFile, TaskContext, TaskResult


UPLOAD_DIR = Path("storage/uploads")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


TEMPLATE_NAME_HINTS = [
    "template",
    "模板",
    "空表",
    "样表",
    "登记表",
    "申请表",
    "审核表",
    "汇总表",
]


class DocumentService:
    def __init__(self, router: DocumentRouter | None = None) -> None:
        self.router = router or DocumentRouter()

    def read(self, file_path: str):
        return self.router.get_reader(file_path).read(file_path)

    def read_text(self, file_path: str) -> str:
        return self.read(file_path).full_text

    def extract_placeholders(self, file_path: str) -> List[str]:
        return self.router.get_inspector(file_path).extract_placeholders(file_path)

    def fill_template(self, template_path: str, data: Dict[str, Any]) -> str:
        return self.router.get_writer(template_path).fill_template(template_path, data)

    def verify(
        self,
        output_path: str,
        expected_placeholders: List[str],
        filled_data: Dict[str, Any],
    ) -> Dict[str, Any]:
        result = self.router.get_verifier(output_path).verify_filled_document(
            output_path=output_path,
            expected_placeholders=expected_placeholders,
            filled_data=filled_data,
        )
        return result


class TaskService:
    def __init__(self):
        from app.core.llm_client import LLMClient

        self.document_service = DocumentService()
        llm_client = LLMClient()
        self.agent = WorkflowAgent(llm=llm_client, document_service=self.document_service)

    def run_workflow_task(
        self,
        files: list[UploadFile],
        roles: list[str],
        user_prompt: str,
    ) -> TaskResult:
        logs = ["service: 开始处理统一工作流请求"]

        try:
            if len(files) != len(roles):
                return TaskResult(
                    success=False,
                    message="文件数量与角色数量不一致",
                    answer="",
                    logs=logs,
                    error="files/roles length mismatch",
                )

            parsed_files: list[InputFile] = []

            for idx, (file, role) in enumerate(zip(files, roles), start=1):
                filename = file.filename or f"file_{idx}.docx"
                file_path = UPLOAD_DIR / f"{idx}_{role}_{filename}"

                with file_path.open("wb") as f:
                    shutil.copyfileobj(file.file, f)

                logs.append(f"service: 文件已保存 -> {file_path}")

                try:
                    parsed_doc = self.document_service.read(str(file_path))
                except Exception as e:
                    logs.append(f"service: 文档解析失败 -> {filename}, detail={str(e)}")
                    return TaskResult(
                        success=False,
                        message="文档解析失败",
                        answer="",
                        logs=logs,
                        error=str(e),
                    )

                logs.append(f"service: docx 解析完成 -> {parsed_doc.file_name} ({role})")

                parsed_files.append(
                    InputFile(
                        role=role,
                        file_name=parsed_doc.file_name,
                        file_path=parsed_doc.file_path,
                        paragraphs=parsed_doc.paragraphs,
                        tables=parsed_doc.tables,
                        full_text=parsed_doc.full_text,
                    )
                )

            source_files = [f for f in parsed_files if f.role == "source"]
            combined_text = "\n\n".join(
                f"===== {f.file_name} =====\n{f.full_text}" for f in source_files
            )

            first_file_name = parsed_files[0].file_name if parsed_files else ""
            first_file_path = parsed_files[0].file_path if parsed_files else ""

            context = TaskContext(
                user_prompt=user_prompt,
                logs=logs,
                files=parsed_files,
                file_name=first_file_name,
                file_path=first_file_path,
                full_text=combined_text,
                paragraphs=[],
                tables=[],
            )

            try:
                agent_files = [
                    {
                        "filename": f.file_name,
                        "path": f.file_path,
                        "role": f.role,
                    }
                    for f in parsed_files
                ]
                agent_result = self.agent.run(
                    user_prompt=user_prompt,
                    files=agent_files,
                )

                output_file_path = self._extract_output_file_path(agent_result)

                result = TaskResult(
                    success=True,
                    message="任务执行成功",
                    answer=agent_result.get("final_answer", ""),
                    logs=logs + [f"agent: {str(agent_result)}"],
                    structured_data=agent_result.get("state", {}),
                    output_file_path=output_file_path,
                )
            except Exception as e:
                logs.append(f"agent error: {str(e)}")
                result = TaskResult(
                    success=False,
                    message="agent执行失败",
                    answer="",
                    logs=logs,
                    error=str(e),
                )

            result.logs.append(f"service: 最终 roles -> {roles}")
            result.logs.append("service: 统一工作流任务处理结束")
            return result

        except Exception as e:
            logs.append(f"service error: {str(e)}")
            return TaskResult(
                success=False,
                message="服务执行失败",
                answer="",
                logs=logs,
                error=str(e),
            )

    def _extract_output_file_path(self, agent_result: Dict[str, Any]) -> str | None:
        state = agent_result.get("state", {}) if isinstance(agent_result, dict) else {}
        if isinstance(state, dict):
            for key in ["filled_doc", "output_file", "filled_doc_result"]:
                value = state.get(key)
                if isinstance(value, dict):
                    output_path = value.get("output_path")
                    if (not isinstance(output_path, str) or not output_path.strip()) and isinstance(value.get("data"), dict):
                        output_path = value["data"].get("output_path")
                    if isinstance(output_path, str) and output_path.strip():
                        return output_path

                if isinstance(value, str) and value.strip().lower().endswith(".docx"):
                    return value

        step_results = agent_result.get("step_results", []) if isinstance(agent_result, dict) else []
        if isinstance(step_results, list):
            for step in reversed(step_results):
                result = step.get("result") if isinstance(step, dict) else None
                if isinstance(result, dict):
                    output_path = result.get("output_path")
                    if (not isinstance(output_path, str) or not output_path.strip()) and isinstance(result.get("data"), dict):
                        output_path = result["data"].get("output_path")
                    if isinstance(output_path, str) and output_path.strip():
                        return output_path

        return None


def infer_roles_from_filenames(files: list[UploadFile]) -> list[str]:
    if not files:
        return []

    if len(files) == 1:
        return ["source"]

    scored = []
    for idx, file in enumerate(files):
        score = _score_template_candidate(file)
        scored.append((idx, score))

    scored.sort(key=lambda x: x[1], reverse=True)
    best_idx, best_score = scored[0]

    if best_score < 2:
        return ["source"] * len(files)

    roles = ["source"] * len(files)
    roles[best_idx] = "template"
    return roles


def _score_template_candidate(file: UploadFile) -> int:
    score = 0
    name = (file.filename or "").lower()

    for hint in TEMPLATE_NAME_HINTS:
        if hint.lower() in name:
            score += 3

    if not name.endswith(".docx"):
        return score

    try:
        current_pos = file.file.tell()
    except Exception:
        current_pos = None

    try:
        doc = Document(file.file)

        placeholder_count = 0
        for para in doc.paragraphs:
            placeholder_count += len(PLACEHOLDER_PATTERN.findall(para.text))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    placeholder_count += len(PLACEHOLDER_PATTERN.findall(cell.text))

        if placeholder_count > 0:
            score += 5

        if len(doc.tables) >= 2:
            score += 1

        empty_cells = 0
        total_cells = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    total_cells += 1
                    if not cell.text.strip():
                        empty_cells += 1

        if total_cells > 0:
            empty_ratio = empty_cells / total_cells
            if empty_ratio >= 0.3:
                score += 2

    except Exception:
        pass
    finally:
        try:
            if current_pos is not None:
                file.file.seek(current_pos)
            else:
                file.file.seek(0)
        except Exception:
            pass

    return score