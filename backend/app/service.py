import re
import shutil
from pathlib import Path

from docx import Document
from fastapi import UploadFile

from app.agent.agent import WorkflowAgent
from app.domain.models import InputFile, TaskContext, TaskResult
from app.tools.word.reader import DocxReader


UPLOAD_DIR = Path("storage/uploads")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


TEMPLATE_NAME_HINTS = [
    "template",
    "模板",
    "空表",
    "样表",
    "登记表",
    "申请表",
    "审核表",
    "汇总表",
]

PLACEHOLDER_PATTERN = re.compile(r"\{\{(.*?)\}\}")


class TaskService:
    def __init__(self):
        self.reader = DocxReader()
        self.agent = WorkflowAgent()

    def run_workflow_task(
        self,
        files: list[UploadFile],
        roles: list[str],
        user_prompt: str,
    ) -> TaskResult:
        logs = ["service: 开始处理统一工作流请求"]

        try:
            if len(files) != len(roles):
                return TaskResult(
                    success=False,
                    message="文件数量与角色数量不一致",
                    answer="",
                    logs=logs,
                    error="files/roles length mismatch",
                )

            parsed_files: list[InputFile] = []

            for idx, (file, role) in enumerate(zip(files, roles), start=1):
                filename = file.filename or f"file_{idx}.docx"
                file_path = UPLOAD_DIR / f"{idx}_{role}_{filename}"

                with file_path.open("wb") as f:
                    shutil.copyfileobj(file.file, f)

                logs.append(f"service: 文件已保存 -> {file_path}")

                parsed_doc = self.reader.read(str(file_path))
                logs.append(f"service: docx 解析完成 -> {parsed_doc.file_name} ({role})")

                parsed_files.append(
                    InputFile(
                        role=role,
                        file_name=parsed_doc.file_name,
                        file_path=parsed_doc.file_path,
                        paragraphs=parsed_doc.paragraphs,
                        tables=parsed_doc.tables,
                        full_text=parsed_doc.full_text,
                    )
                )

            source_files = [f for f in parsed_files if f.role == "source"]
            combined_text = "\n\n".join(
                f"===== {f.file_name} =====\n{f.full_text}" for f in source_files
            )

            first_file_name = parsed_files[0].file_name if parsed_files else ""
            first_file_path = parsed_files[0].file_path if parsed_files else ""

            context = TaskContext(
                user_prompt=user_prompt,
                logs=logs,
                files=parsed_files,
                file_name=first_file_name,
                file_path=first_file_path,
                full_text=combined_text,
                paragraphs=[],
                tables=[],
            )

            result = self.agent.run(context)
            result.logs.append(f"service: 最终 roles -> {roles}")
            result.logs.append("service: 统一工作流任务处理结束")
            return result

        except Exception as e:
            logs.append(f"service error: {str(e)}")
            return TaskResult(
                success=False,
                message="服务执行失败",
                answer="",
                logs=logs,
                error=str(e),
            )


def infer_roles_from_filenames(files: list[UploadFile]) -> list[str]:
    """
    自动推断文件角色：
    - 优先找 template
    - 其余默认为 source
    """
    if not files:
        return []

    if len(files) == 1:
        return ["source"]

    scored = []
    for idx, file in enumerate(files):
        score = _score_template_candidate(file)
        scored.append((idx, score))

    # 找最像 template 的文件
    scored.sort(key=lambda x: x[1], reverse=True)
    best_idx, best_score = scored[0]

    # 如果最高分太低，说明都不像模板 -> 全部 source
    if best_score < 2:
        return ["source"] * len(files)

    roles = ["source"] * len(files)
    roles[best_idx] = "template"
    return roles


def _score_template_candidate(file: UploadFile) -> int:
    score = 0
    name = (file.filename or "").lower()

    # 1. 文件名规则
    for hint in TEMPLATE_NAME_HINTS:
        if hint.lower() in name:
            score += 3

    # 2. 只对 docx 做内容规则
    if not name.endswith(".docx"):
        return score

    # 保存当前文件指针位置，避免影响后续读取
    try:
        current_pos = file.file.tell()
    except Exception:
        current_pos = None

    try:
        doc = Document(file.file)

        # 2.1 占位符
        placeholder_count = 0
        for para in doc.paragraphs:
            placeholder_count += len(PLACEHOLDER_PATTERN.findall(para.text))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    placeholder_count += len(PLACEHOLDER_PATTERN.findall(cell.text))

        if placeholder_count > 0:
            score += 5

        # 2.2 表格较多
        if len(doc.tables) >= 2:
            score += 1

        # 2.3 空白单元格较多
        empty_cells = 0
        total_cells = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    total_cells += 1
                    if not cell.text.strip():
                        empty_cells += 1

        if total_cells > 0:
            empty_ratio = empty_cells / total_cells
            if empty_ratio >= 0.3:
                score += 2

    except Exception:
        # 内容分析失败就只用文件名分
        pass
    finally:
        try:
            if current_pos is not None:
                file.file.seek(current_pos)
            else:
                file.file.seek(0)
        except Exception:
            pass

    return score