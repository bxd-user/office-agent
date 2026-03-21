from docx import Document

from app.executor import StepExecutor
from app.models import TaskContext, TaskResult
from app.planner import WorkflowPlanner
from app.prompts import (
    SUMMARIZE_SYSTEM_PROMPT,
    EXTRACT_SYSTEM_PROMPT,
    FILL_SYSTEM_PROMPT,
)
from app.tools.docx_writer import DocxWriter
from app.tools.llm_client import LLMClient
from app.field_config import EXTRACT_FIELD_CANDIDATES


class WorkflowAgent:
    def __init__(self):
        self.llm = LLMClient()
        self.writer = DocxWriter()
        self.planner = WorkflowPlanner()
        self.executor = StepExecutor()

        self.SUMMARIZE_SYSTEM_PROMPT = SUMMARIZE_SYSTEM_PROMPT
        self.EXTRACT_SYSTEM_PROMPT = EXTRACT_SYSTEM_PROMPT
        self.FILL_SYSTEM_PROMPT = FILL_SYSTEM_PROMPT

    def run(self, context: TaskContext) -> TaskResult:
        logs = list(context.logs)
        logs.append("agent: 开始执行")

        try:
            file_roles = [file.role for file in context.files]
            plan = self.planner.create_plan(context.user_prompt, file_roles)
            logs.append(f"agent: 规划任务类型 -> {plan.task_type}")
            logs.append(f"agent: 执行步骤 -> {plan.steps}")

            return self.executor.run(
                agent=self,
                context=context,
                task_type=plan.task_type,
                steps=plan.steps,
                logs=logs,
            )

        except Exception as e:
            logs.append(f"agent error: {str(e)}")
            return TaskResult(
                success=False,
                message="执行失败",
                answer="",
                logs=logs,
                error=str(e),
            )

    def _extract_requested_fields(self, prompt: str) -> list[str]:
        return [field for field in EXTRACT_FIELD_CANDIDATES if field in prompt]

    def _extract_template_placeholders(self, context: TaskContext) -> list[str]:
        template_file = None
        for file in context.files:
            if file.role == "template":
                template_file = file
                break

        if not template_file:
            return []

        import re

        pattern = re.compile(r"\{\{(.*?)\}\}")
        doc = Document(template_file.file_path)

        found = []

        for para in doc.paragraphs:
            found.extend(pattern.findall(para.text))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    found.extend(pattern.findall(cell.text))

        result = []
        seen = set()
        for item in found:
            key = item.strip()
            if key and key not in seen:
                seen.add(key)
                result.append(key)

        return result