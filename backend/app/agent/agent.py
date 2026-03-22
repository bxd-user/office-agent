from docx import Document

from app.agent.executor import StepExecutor
from app.domain.models import TaskContext, TaskResult
from app.agent.planner import WorkflowPlanner, PlannerContext, PlannerFileContext
from app.agent.prompts import (
    SUMMARIZE_SYSTEM_PROMPT,
    EXTRACT_SYSTEM_PROMPT,
    FILL_SYSTEM_PROMPT,
)
from app.tools.word.writer import DocxWriter
from app.core.llm_client import LLMClient
from app.domain.field_config import EXTRACT_FIELD_CANDIDATES

class WorkflowAgent:
    def __init__(self):
        self.llm = LLMClient()
        self.writer = DocxWriter()
        self.planner = WorkflowPlanner(self.llm)
        self.executor = StepExecutor()

        self.SUMMARIZE_SYSTEM_PROMPT = SUMMARIZE_SYSTEM_PROMPT
        self.EXTRACT_SYSTEM_PROMPT = EXTRACT_SYSTEM_PROMPT
        self.FILL_SYSTEM_PROMPT = FILL_SYSTEM_PROMPT

    def run(self, context: TaskContext) -> TaskResult:
        logs = list(context.logs)
        logs.append("agent: 开始执行")

        try:
            planner_context = PlannerContext(
                user_prompt=context.user_prompt,
                files=[
                    PlannerFileContext(
                        file_id=f"file_{i+1}",
                        filename=f.file_name,
                        file_type="word",
                        extension=".docx",
                        role_hint=f.role,
                        content_preview=f.full_text[:500] if f.full_text else "",
                    )
                    for i, f in enumerate(context.files)
                ]
            )

            plan = self.planner.create_plan(planner_context)
            logs.append(f"agent: 规划任务类型 -> {plan.task_type}")
            logs.append(f"agent: 执行步骤数量 -> {len(plan.steps)}")

            self._apply_plan_file_roles(context, plan.file_roles, logs)
            self._refresh_context_legacy_fields(context)

            return self.executor.run(
                agent=self,
                context=context,
                plan=plan,
                logs=logs,
            )

        except Exception as e:
            logs.append(f"agent error: {str(e)}")
            return TaskResult(
                success=False,
                message="执行失败",
                answer="",
                logs=logs,
                error=str(e),
            )

    def _extract_requested_fields(self, prompt: str) -> list[str]:
        return [field for field in EXTRACT_FIELD_CANDIDATES if field in prompt]

    def _apply_plan_file_roles(self, context: TaskContext, file_roles: dict[str, str], logs: list[str]) -> None:
        if not file_roles:
            logs.append("agent: planner 未返回 file_roles，沿用输入角色")
            return

        for i, file in enumerate(context.files, start=1):
            file_id = f"file_{i}"
            new_role = file_roles.get(file_id)
            if new_role and new_role != file.role:
                logs.append(f"agent: 角色修正 {file.file_name}: {file.role} -> {new_role}")
                file.role = new_role

    def _refresh_context_legacy_fields(self, context: TaskContext) -> None:
        source_files = [f for f in context.files if f.role == "source"]
        template_file = next((f for f in context.files if f.role == "template"), None)

        if source_files:
            context.file_name = source_files[0].file_name
            context.file_path = source_files[0].file_path
            context.full_text = "\n\n".join(
                f"===== {f.file_name} =====\n{f.full_text}" for f in source_files
            )

        if template_file:
            context.template_file_name = template_file.file_name
            context.template_file_path = template_file.file_path

    def _extract_template_placeholders(self, context: TaskContext) -> list[str]:
        template_file = None
        for file in context.files:
            if file.role == "template":
                template_file = file
                break

        if not template_file:
            return []

        import re

        pattern = re.compile(r"\{\{(.*?)\}\}")
        doc = Document(template_file.file_path)

        found = []

        for para in doc.paragraphs:
            found.extend(pattern.findall(para.text))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    found.extend(pattern.findall(cell.text))

        result = []
        seen = set()
        for item in found:
            key = item.strip()
            if key and key not in seen:
                seen.add(key)
                result.append(key)

        return result
