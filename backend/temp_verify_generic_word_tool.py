from pathlib import Path
import sys

sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx import Document
from app.tools.word_tool import WordTool

base = Path("storage/outputs")
base.mkdir(parents=True, exist_ok=True)
src = base / "generic_word_src.docx"
out = base / "generic_word_out.docx"

# 构造通用文档：文章 + 表格
doc = Document()
doc.add_heading("通用文档测试", level=1)
doc.add_paragraph("这是正文内容，作者：{{作者}}")
doc.add_paragraph("摘要：${摘要}")

table = doc.add_table(rows=2, cols=2)
table.style = "Table Grid"
table.cell(0, 0).text = "字段"
table.cell(0, 1).text = "值"
table.cell(1, 0).text = "城市："
table.cell(1, 1).text = ""

doc.save(src)

wt = WordTool()
payload = wt.read_document_for_llm(str(src))
print("fields:", payload.get("fields", []))

wt.fill_document_with_context(
    str(src),
    {"作者": "Alice", "摘要": "用于测试通用word工具", "城市": "北京"},
    str(out),
)

verify = wt.read_document_for_llm(str(out))
print("p1:", verify["content"]["body"]["paragraphs"][1]["text"])
print("p2:", verify["content"]["body"]["paragraphs"][2]["text"])
print("table:", verify["content"]["body"]["tables"][0]["rows"][1][1])
