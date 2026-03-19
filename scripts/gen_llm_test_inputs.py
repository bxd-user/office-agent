from __future__ import annotations

from openpyxl import Workbook
from docx import Document


def main() -> None:
    xlsx_path = r"e:\office-agent\storage\uploads\llm_test_input.xlsx"
    docx_path = r"e:\office-agent\storage\uploads\llm_test_template.docx"
    semantic_docx_path = r"e:\office-agent\storage\uploads\llm_test_template_semantic.docx"

    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Sheet1"
    worksheet.append(["姓名", "班级", "学号"])
    worksheet.append(["张三", "三年二班", "2026001"])
    workbook.save(xlsx_path)

    document = Document()
    document.add_paragraph("学生信息")
    document.add_paragraph("姓名：{{姓名}}")
    document.add_paragraph("班级：{{班级}}")
    document.add_paragraph("学号：{{学号}}")
    document.save(docx_path)

    semantic_document = Document()
    semantic_document.add_paragraph("学生信息")
    semantic_document.add_paragraph("学生姓名：{{学生姓名}}")
    semantic_document.add_paragraph("所在班级：{{所在班级}}")
    semantic_document.add_paragraph("学生编号：{{学生编号}}")
    semantic_document.save(semantic_docx_path)

    print(xlsx_path)
    print(docx_path)
    print(semantic_docx_path)


if __name__ == "__main__":
    main()